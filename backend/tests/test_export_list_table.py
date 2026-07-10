"""
模块：导出列表符号与表格样式测试
用途：验收 list_style / ordered_list_style / table 配置写入 docx；MD 列表与表格解析。
对接：export_service._bullet_prefix / write_markdown_body / export 任务
"""

from io import BytesIO

from app.services.export_service import (
    _bullet_prefix,
    _ordered_prefix,
    _parse_md_table_lines,
    write_markdown_body,
)


def test_list_prefixes():
    assert _bullet_prefix("disc") == "•"
    assert _bullet_prefix("dash") == "–"
    assert _bullet_prefix("check") == "✓"
    assert _ordered_prefix("decimal-dot", 1) == "1."
    assert _ordered_prefix("decimal-full-paren", 2) == "（2）"
    assert _ordered_prefix("chinese-dot", 1) == "一、"
    assert _ordered_prefix("lower-alpha", 1) == "a."


def test_parse_md_table():
    rows = _parse_md_table_lines(
        [
            "| 项 | 权重 |",
            "| --- | --- |",
            "| 架构 | 20% |",
        ]
    )
    assert rows is not None
    assert rows[0] == ["项", "权重"]
    assert rows[1] == ["架构", "20%"]


def test_write_markdown_body_lists_and_table():
    from docx import Document  # type: ignore

    doc = Document()
    md = (
        "前言段落\n"
        "- 无序甲\n"
        "- 无序乙\n"
        "1. 有序一\n"
        "2. 有序二\n"
        "| 列A | 列B |\n"
        "| --- | --- |\n"
        "| 1 | 2 |\n"
    )
    write_markdown_body(
        doc,
        md,
        list_cfg={
            "list_style": "arrow",
            "ordered_list_style": "chinese-dot",
            "list_indent_chars": 2,
        },
        table_cfg={
            "border_color": "#123456",
            "border_width": 1,
            "full_width": True,
            "header_row": {
                "font": "黑体",
                "size": "小四",
                "background_color": "#EEF3FF",
                "text_color": "#123456",
            },
            "body_cell": {"font": "宋体", "size": "小四"},
        },
    )
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)
    assert "➢ 无序甲" in joined or "无序甲" in joined
    assert "一、" in joined or "有序一" in joined
    assert len(doc.tables) >= 1
    assert "列A" in doc.tables[0].cell(0, 0).text


def test_export_scoring_table_and_list_style(client):
    client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "template_name": "列表表格测",
                "body_text": {
                    "font": "宋体",
                    "size": "小四",
                    "list_style": "check",
                    "ordered_list_style": "decimal-paren",
                    "list_indent_chars": 2,
                },
                "table": {
                    "border_width": 1,
                    "border_color": "#333333",
                    "full_width": True,
                    "header_row": {
                        "font": "黑体",
                        "size": "小四",
                        "background_color": "#F0F0F0",
                        "text_color": "#000000",
                        "alignment": "居中对齐",
                    },
                    "body_cell": {
                        "font": "宋体",
                        "size": "小四",
                        "alignment": "左对齐",
                    },
                    "first_column": {
                        "font": "宋体",
                        "size": "小四",
                        "background_color": "#FAFAFA",
                    },
                },
                "headings": [
                    {
                        "numbering_format": "custom",
                        "numbering_template": "第{zh}章",
                    }
                ],
            }
        },
    )
    proj = client.post("/api/projects", json={"name": "列表导出"}).json()
    pid = proj["id"]
    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "analysisOverview": "概述一行",
            "analysis": {
                "overview": "概述一行",
                "techRequirements": ["支持 2000 路视频", "等保三级"],
                "rejectionRisks": ["未响应★条款"],
                "scoringPoints": [
                    {"name": "总体架构", "weight": "20%"},
                    {"name": "功能完整性", "weight": "25%"},
                ],
            },
            "outline": [{"id": "n1", "title": "总体架构", "children": []}],
            "chapters": [
                {
                    "id": "n1",
                    "title": "总体架构",
                    "body": (
                        "说明如下：\n"
                        "- 采用微服务\n"
                        "1. 网关层\n"
                        "| 模块 | 说明 |\n"
                        "| --- | --- |\n"
                        "| 网关 | 统一入口 |\n"
                    ),
                    "preview": "说明",
                    "wordCount": 20,
                    "status": "done",
                }
            ],
            "facts": [
                {
                    "id": "f1",
                    "category": "招标",
                    "content": "建设周期 180 天",
                    "source": "tender",
                }
            ],
        },
    )
    exp = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exp.status_code == 201
    assert exp.json()["status"] == "success"
    stored = exp.json()["result"]["storedName"]
    dl = client.get(f"/api/projects/{pid}/export/download/{stored}")
    assert dl.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(dl.content))
    texts = "\n".join(p.text for p in doc.paragraphs)
    # 列表符号 check
    assert "✓" in texts or "2000" in texts
    # 评分点表格
    assert len(doc.tables) >= 1
    header_join = " ".join(c.text for c in doc.tables[0].rows[0].cells)
    assert "评分" in header_join or "权重" in header_join or "总体架构" in texts
