"""
模块：V1-D 合成文档结构质量门（failure-first）
用途：用系统 TEMP 合成 DOCX/PDF/TXT/MD 锁定 lightweight 标题/表格/页序与扫描提示契约。
对接：app.services.parse_service.parse_file_to_markdown；V1-D 契约 §3–§7。
二次开发：仅合成数据与标准库/既有依赖；禁止 skip/xfail、fake engine、真实业务样本与数据库。
"""

from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from app.services.parse_service import parse_file_to_markdown

# ---------------------------------------------------------------------------
# 集中 ASCII 锚点（合成样本唯一真值，禁止散落魔法字符串）
# ---------------------------------------------------------------------------
ANCHOR_HEADING = "HEADING1_SYNTH"
# Heading 1–6 唯一 ASCII 锚点（与单元格 H1_SYNTH/H2_SYNTH 刻意区分）
ANCHOR_HEADINGS = (
    "HEADING1_SYNTH",
    "HEADING2_SYNTH",
    "HEADING3_SYNTH",
    "HEADING4_SYNTH",
    "HEADING5_SYNTH",
    "HEADING6_SYNTH",
)
ANCHOR_BEFORE = "BEFORE_TABLE_SYNTH"
ANCHOR_AFTER = "AFTER_TABLE_SYNTH"
ANCHOR_H1 = "H1_SYNTH"
ANCHOR_H2 = "H2_SYNTH"
ANCHOR_A1 = "A1_SYNTH"
ANCHOR_A2 = "A2_SYNTH"
ANCHOR_COL_A = "COLA_SYNTH"
ANCHOR_COL_B = "COLB_SYNTH"
ANCHOR_PIPE = "V|pipe"
ANCHOR_PIPE_ESCAPED = r"V\|pipe"
ANCHOR_MULTI_RAW_LINES = ("LINE1_SYNTH", "LINE2_SYNTH")
ANCHOR_MULTI_FOLDED = "LINE1_SYNTH LINE2_SYNTH"
ANCHOR_CELL_OK = "OK_SYNTH"
ANCHOR_PAGE1 = "PAGE1_SYNTH"
ANCHOR_PAGE2 = "PAGE2_SYNTH"
ANCHOR_TXT = "TXT_BODY_SYNTH"
ANCHOR_MD = "MD_BODY_SYNTH"
ANCHOR_EMPTY_TBL_BEFORE = "EMPTY_TBL_BEFORE_SYNTH"
ANCHOR_EMPTY_TBL_AFTER = "EMPTY_TBL_AFTER_SYNTH"

HEADER_PREFIX = "# 解析结果："
EMPTY_DOCX_PLACEHOLDER = "（DOCX 无段落文本）"
EMPTY_TABLE_PLACEHOLDER = "（空表）"
SCAN_HINT = "（本页未提取到文本，可能是扫描件，请用本地 MinerU）"
GFM_SEP_2 = "| --- | --- |"

# 本模块创建的全部 TEMP 根；清理后必须为空存在性，且不得用空列表假绿
_TEMP_ROOTS: list[Path] = []


def _track_temp_root(root: Path) -> Path:
    """用途：显式登记 TEMP 根，供成功/失败路径统一断言已删除。"""
    resolved = root.resolve()
    _TEMP_ROOTS.append(resolved)
    return resolved


def _make_temp_dir() -> Path:
    """用途：创建并登记专用 TEMP 目录。"""
    return _track_temp_root(Path(tempfile.mkdtemp(prefix="v1d_synth_")))


def _cleanup_temp_root(root: Path) -> None:
    """用途：递归删除 TEMP 根；忽略已不存在。"""
    if root.exists():
        shutil.rmtree(root, ignore_errors=False)
    if root.exists():
        shutil.rmtree(root, ignore_errors=True)


def _assert_temp_gone(root: Path) -> None:
    """用途：失败即红，禁止 TEMP 泄漏。"""
    assert not root.exists(), f"TEMP 根仍存在: {root}"


def _pdf_escape_literal(text: str) -> str:
    """用途：PDF 字符串字面量转义（标准库构造，不用 reportlab）。"""
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_two_page_text_pdf(page1: str, page2: str) -> bytes:
    """
    用途：仅用标准库字节构造可被 pypdf 提取的两页文字 PDF。
    二次开发：禁止 reportlab 或其它 PDF 写库。
    """
    streams = []
    for text in (page1, page2):
        content = f"BT /F1 12 Tf 72 720 Td ({_pdf_escape_literal(text)}) Tj ET"
        streams.append(content.encode("latin-1"))

    objs: list[bytes] = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >>\nendobj\n",
        (
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 5 0 R /Resources << /Font << /F1 7 0 R >> >> >>\nendobj\n"
        ),
        (
            b"4 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 6 0 R /Resources << /Font << /F1 7 0 R >> >> >>\nendobj\n"
        ),
        (
            f"5 0 obj\n<< /Length {len(streams[0])} >>\nstream\n".encode("latin-1")
            + streams[0]
            + b"\nendstream\nendobj\n"
        ),
        (
            f"6 0 obj\n<< /Length {len(streams[1])} >>\nstream\n".encode("latin-1")
            + streams[1]
            + b"\nendstream\nendobj\n"
        ),
        b"7 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objs:
        offsets.append(len(out))
        out.extend(obj)
    xref_at = len(out)
    out.extend(f"xref\n0 {len(objs) + 1}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets:
        out.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        (
            f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n"
        ).encode("latin-1")
    )
    return bytes(out)


def _write_heading_table_docx(path: Path) -> None:
    """用途：Heading1–6 → 表前段 → 2x2 普通表 → 表后段。"""
    doc = Document()
    for level, anchor in enumerate(ANCHOR_HEADINGS, start=1):
        doc.add_heading(anchor, level=level)
    doc.add_paragraph(ANCHOR_BEFORE)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = ANCHOR_H1
    table.cell(0, 1).text = ANCHOR_H2
    table.cell(1, 0).text = ANCHOR_A1
    table.cell(1, 1).text = ANCHOR_A2
    doc.add_paragraph(ANCHOR_AFTER)
    doc.save(str(path))


def _write_empty_table_docx(path: Path) -> None:
    """用途：唯一表前段 → rows=0/cols=2 空表 → 唯一表后段。"""
    doc = Document()
    doc.add_paragraph(ANCHOR_EMPTY_TBL_BEFORE)
    doc.add_table(rows=0, cols=2)
    doc.add_paragraph(ANCHOR_EMPTY_TBL_AFTER)
    doc.save(str(path))


def _write_special_cells_docx(path: Path) -> None:
    """用途：空单元格、管道符、单元格内换行。"""
    doc = Document()
    doc.add_paragraph("SPECIAL_TABLE_INTRO")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = ANCHOR_COL_A
    table.cell(0, 1).text = ANCHOR_COL_B
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ANCHOR_PIPE
    # 单元格换行：两段落入同一单元格
    multi = table.cell(2, 0)
    multi.text = ""
    multi.paragraphs[0].add_run(ANCHOR_MULTI_RAW_LINES[0])
    multi.add_paragraph(ANCHOR_MULTI_RAW_LINES[1])
    table.cell(2, 1).text = ANCHOR_CELL_OK
    doc.save(str(path))


def _write_empty_docx(path: Path) -> None:
    """用途：仅空段落，保持既有占位文案。"""
    doc = Document()
    # 默认已有空段；再显式加空段，确保无可见文本
    doc.add_paragraph("")
    doc.save(str(path))


def test_docx_heading_before_table_after_block_order_and_gfm():
    """DOCX：Heading1–6/表前/2x2 表/表后 — 精确整行集合、级别计数与完整输出锁定。"""
    root = _make_temp_dir()
    try:
        name = "synth_heading_table.docx"
        path = root / name
        _write_heading_table_docx(path)
        assert path.is_file() and path.stat().st_size > 0

        md = parse_file_to_markdown(path, name)

        # Heading 1–6 → Markdown # 至 ######（精确整行 + 单空格 + 锚点）
        heading_lines = [
            f"{'#' * level} {anchor}" for level, anchor in enumerate(ANCHOR_HEADINGS, start=1)
        ]
        header_row = f"| {ANCHOR_H1} | {ANCHOR_H2} |"
        data_row = f"| {ANCHOR_A1} | {ANCHOR_A2} |"
        # 完整输出锁定：header、六标题、before、GFM 三行、after、块间空行与尾换行
        expected = (
            f"{HEADER_PREFIX}{name}\n\n"
            + "\n\n".join(heading_lines)
            + f"\n\n{ANCHOR_BEFORE}\n\n"
            + f"{header_row}\n{GFM_SEP_2}\n{data_row}\n\n"
            + f"{ANCHOR_AFTER}\n"
        )
        lines = md.splitlines()

        # 六级诊断：每个精确标题行在 splitlines 整行集合中恰好出现一次
        for level, expected_line in enumerate(heading_lines, start=1):
            count = lines.count(expected_line)
            assert count == 1, (
                f"Heading {level} 精确行出现次数异常: expect=1 got={count} line={expected_line!r}"
            )

        # 表前、GFM 三行、表后：精确整行计数
        for label, expected_line in (
            ("before", ANCHOR_BEFORE),
            ("gfm_header", header_row),
            ("gfm_sep", GFM_SEP_2),
            ("gfm_data", data_row),
            ("after", ANCHOR_AFTER),
        ):
            count = lines.count(expected_line)
            assert count == 1, f"{label} 精确行出现次数异常: expect=1 got={count} line={expected_line!r}"

        # 顺序：精确行的行号（禁止 substring index 证明级别）
        def _line_no(exact: str) -> int:
            return lines.index(exact)

        h_nos = [_line_no(hl) for hl in heading_lines]
        for left, right in zip(h_nos, h_nos[1:]):
            assert left < right
        assert (
            h_nos[-1]
            < _line_no(ANCHOR_BEFORE)
            < _line_no(header_row)
            < _line_no(GFM_SEP_2)
            < _line_no(data_row)
            < _line_no(ANCHOR_AFTER)
        )

        # 完整相等同时关闭额外乱码/错误间距
        assert md == expected
        # 兼容旧常量：H1 锚点仍为 ANCHOR_HEADING
        assert ANCHOR_HEADING == ANCHOR_HEADINGS[0]
        assert heading_lines[0] == f"# {ANCHOR_HEADING}"
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_docx_empty_cell_pipe_escape_and_newline_fold():
    """DOCX：空单元格、V|pipe 转义、单元格换行折叠为单空格。"""
    root = _make_temp_dir()
    try:
        name = "synth_special_cells.docx"
        path = root / name
        _write_special_cells_docx(path)
        assert path.is_file() and path.stat().st_size > 0

        md = parse_file_to_markdown(path, name)
        assert md.startswith(f"{HEADER_PREFIX}{name}\n\n")

        header_row = f"| {ANCHOR_COL_A} | {ANCHOR_COL_B} |"
        assert header_row in md
        assert GFM_SEP_2 in md

        # 空单元格稳定列位：|  |
        empty_and_pipe_row = f"|  | {ANCHOR_PIPE_ESCAPED} |"
        assert "|  |" in md
        assert ANCHOR_PIPE_ESCAPED in md
        assert empty_and_pipe_row in md
        # 原始未转义管道不得以裸表单元形式出现
        assert f"| {ANCHOR_PIPE} |" not in md

        # 换行折叠为单空格
        folded_row = f"| {ANCHOR_MULTI_FOLDED} | {ANCHOR_CELL_OK} |"
        assert ANCHOR_MULTI_FOLDED in md
        assert folded_row in md
        assert f"{ANCHOR_MULTI_RAW_LINES[0]}\n{ANCHOR_MULTI_RAW_LINES[1]}" not in md

        i_header = md.index(header_row)
        i_empty = md.index(empty_and_pipe_row)
        i_folded = md.index(folded_row)
        assert i_header < i_empty < i_folded
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_docx_empty_keeps_placeholder():
    """空 DOCX 保持既有「（DOCX 无段落文本）」。"""
    root = _make_temp_dir()
    try:
        name = "synth_empty.docx"
        path = root / name
        _write_empty_docx(path)
        md = parse_file_to_markdown(path, name)
        assert md == f"{HEADER_PREFIX}{name}\n\n{EMPTY_DOCX_PLACEHOLDER}\n"
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_docx_zero_row_table_empty_placeholder_block_order():
    """DOCX：rows=0/cols=2 空表完整相等固定块，禁止重复或额外内容。"""
    root = _make_temp_dir()
    try:
        name = "synth_empty_table.docx"
        path = root / name
        _write_empty_table_docx(path)
        assert path.is_file() and path.stat().st_size > 0

        md = parse_file_to_markdown(path, name)
        # 固定：header + before + 两换行 + （空表） + 两换行 + after + 尾换行
        expected = (
            f"{HEADER_PREFIX}{name}\n\n"
            f"{ANCHOR_EMPTY_TBL_BEFORE}\n\n"
            f"{EMPTY_TABLE_PLACEHOLDER}\n\n"
            f"{ANCHOR_EMPTY_TBL_AFTER}\n"
        )
        assert md == expected
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_txt_and_md_representative_regression():
    """TXT/MD 代表样本：头部与正文回归，协议不变。"""
    root = _make_temp_dir()
    try:
        txt_name = "synth_body.txt"
        txt_path = root / txt_name
        txt_path.write_text(f"{ANCHOR_TXT}\n", encoding="utf-8")
        txt_md = parse_file_to_markdown(txt_path, txt_name)
        assert txt_md == f"{HEADER_PREFIX}{txt_name}\n\n{ANCHOR_TXT}\n"

        md_name = "synth_body.md"
        md_path = root / md_name
        md_path.write_text(f"# Title\n\n{ANCHOR_MD}\n", encoding="utf-8")
        out = parse_file_to_markdown(md_path, md_name)
        assert out.startswith(f"{HEADER_PREFIX}{md_name}\n\n")
        assert ANCHOR_MD in out
        assert out == f"{HEADER_PREFIX}{md_name}\n\n# Title\n\n{ANCHOR_MD}\n"
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_pdf_two_page_text_order_without_scan_hint():
    """两页文字 PDF：页标题与 PAGE1/PAGE2 严格顺序，且无扫描提示。"""
    root = _make_temp_dir()
    try:
        name = "synth_two_page_text.pdf"
        path = root / name
        path.write_bytes(build_two_page_text_pdf(ANCHOR_PAGE1, ANCHOR_PAGE2))
        assert path.is_file() and path.stat().st_size > 0

        md = parse_file_to_markdown(path, name)
        assert md.startswith(f"{HEADER_PREFIX}{name}\n\n")

        p1 = "## 第 1 页"
        p2 = "## 第 2 页"
        assert p1 in md
        assert p2 in md
        assert ANCHOR_PAGE1 in md
        assert ANCHOR_PAGE2 in md
        assert SCAN_HINT not in md
        assert md.count(SCAN_HINT) == 0

        i_p1 = md.index(p1)
        i_a1 = md.index(ANCHOR_PAGE1)
        i_p2 = md.index(p2)
        i_a2 = md.index(ANCHOR_PAGE2)
        assert i_p1 < i_a1 < i_p2 < i_a2
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_pdf_two_blank_pages_scan_hint_count_exactly_two():
    """PdfWriter 两页 blank：扫描提示 count=2 且严格 p1<hint1<p2<hint2。"""
    root = _make_temp_dir()
    try:
        name = "synth_two_blank.pdf"
        path = root / name
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_blank_page(width=612, height=792)
        with path.open("wb") as fh:
            writer.write(fh)
        assert path.is_file() and path.stat().st_size > 0

        md = parse_file_to_markdown(path, name)
        assert md.startswith(f"{HEADER_PREFIX}{name}\n\n")
        p1 = "## 第 1 页"
        p2 = "## 第 2 页"
        assert p1 in md
        assert p2 in md
        assert md.count(SCAN_HINT) == 2
        # 严格页内顺序：第1页标题 < 第1页提示 < 第2页标题 < 第2页提示
        i_p1 = md.index(p1)
        i_hint1 = md.index(SCAN_HINT)
        i_p2 = md.index(p2)
        i_hint2 = md.index(SCAN_HINT, i_hint1 + 1)
        assert i_p1 < i_hint1 < i_p2 < i_hint2
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


def test_temp_roots_explicitly_tracked_and_absent_after_tests():
    """
    TEMP 根必须显式记录且清理后不存在（本用例自建 TEMP，不依赖其它测试先运行）。
    禁止用空列表 for 循环制造假绿：登记表在本会话中须非空。
    """
    before_len = len(_TEMP_ROOTS)
    root = _make_temp_dir()
    try:
        # 本次登记增量、绝对路径、前缀
        assert len(_TEMP_ROOTS) == before_len + 1, "本次 TEMP 登记增量必须为 +1"
        assert _TEMP_ROOTS[-1] == root.resolve()
        assert isinstance(root, Path)
        assert root.is_absolute()
        assert "v1d_synth_" in root.name
        assert root.exists() and root.is_dir()
        marker = root / "temp_track_marker.txt"
        marker.write_text("TEMP_TRACK_SYNTH\n", encoding="utf-8")
        assert marker.is_file() and marker.stat().st_size > 0

        assert len(_TEMP_ROOTS) > 0, "TEMP 根登记表为空，疑似未创建真实样本"
        for p in _TEMP_ROOTS:
            assert isinstance(p, Path)
            assert p.is_absolute()
            assert "v1d_synth_" in p.name
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)

    still_alive = [str(p) for p in _TEMP_ROOTS if p.exists()]
    assert still_alive == [], f"仍有未清理 TEMP: {still_alive}"


def test_source_anti_fake_green_guards():
    """测试源码反假绿：禁止 skip/xfail/importorskip、宽松 or 短路、fake engine、真实数据路径。"""
    src_path = Path(__file__).resolve()
    src = src_path.read_text(encoding="utf-8")

    # 排除本守卫函数，避免拼接片段自匹配；再去掉文档字符串
    guard_name = "test_source_anti_fake_green_guards"
    g0 = src.index("def " + guard_name)
    g1_match = re.search(r"\ndef\s+\w+", src[g0 + 4 :])
    g1 = (g0 + 4 + g1_match.start()) if g1_match else len(src)
    scannable = src[:g0] + src[g1:]
    code_only = re.sub(r'"""[\s\S]*?"""', "", scannable)
    code_only = re.sub(r"'''[\s\S]*?'''", "", code_only)

    def _j(*parts: str) -> str:
        return "".join(parts)

    banned_callables = [
        _j("pytest", ".skip("),
        _j("pytest", ".xfail("),
        _j("pytest", ".importorskip("),
        _j("importorskip", "("),
        _j("@pytest.mark.", "skip"),
        _j("@pytest.mark.", "xfail"),
        _j("unittest", ".skip"),
        _j("register", "_engine("),
        _j("import ", "reportlab"),
        _j("from ", "reportlab"),
        _j("Fake", "ParseEngine"),
        _j("biaoshu", ".db"),
        _j("API", "_KEY"),
    ]
    for token in banned_callables:
        assert token not in code_only, f"反假绿命中禁止片段: {token}"

    # 禁止指向真实 uploads / 业务目录的硬编码
    assert _j("biaoshu", "\\uploads") not in code_only
    assert _j("biaoshu", "/uploads") not in code_only
    assert re.search(r"\bCookie\b", code_only) is None

    # 禁止宽松 or 断言假绿
    loose_or = []
    for line in code_only.splitlines():
        stripped = line.strip()
        if stripped.startswith("assert ") and re.search(r"\bor\b", stripped):
            loose_or.append(stripped)
    assert loose_or == [], f"发现宽松 or 断言: {loose_or}"

    # 不得对空列表做 for 断言假绿
    empty_for = re.findall(r"for\s+\w+\s+in\s+\[\s*\]\s*:", code_only)
    assert empty_for == [], f"发现空列表 for 假绿: {empty_for}"

    # 必须真实调用生产解析函数
    assert "parse_file_to_markdown" in code_only
    assert "from app.services.parse_service import parse_file_to_markdown" in code_only
