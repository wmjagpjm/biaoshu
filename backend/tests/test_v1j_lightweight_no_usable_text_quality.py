"""
模块：V1-J lightweight 无有效正文质量门 failure-first 专项
用途：锁定空 TXT/MD/DOCX/PDF、项目 parse 零副作用与知识库清 chunk 契约；
  生产未改时必须业务断言红，禁止 import/fixture/环境红、skip/xfail 或条件假绿。
对接：parse_service.parse_file_to_markdown；task_service parse；knowledge_service 索引。
二次开发：仅系统 TEMP 合成样本与 pytest 临时 upload 根；禁止真实 DB/uploads/联网/模型。
"""

from __future__ import annotations

import copy
import json
import re
import shutil
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from sqlalchemy import func, select

from app.core.config import get_settings
from app.core.database import SessionLocal
from app.models.entities import KbChunkRow, KbDocumentRow
from app.services import parse_service
from app.services.parse_service import parse_file_to_markdown

# 契约 §2.2 固定中文错误（精确相等；match 无正则特殊字符）
FIXED_ERR = "未提取到可用正文，请检查文件是否为空；扫描版 PDF 请使用本地 MinerU"
HEADER_PREFIX = "# 解析结果："
EMPTY_TABLE_PLACEHOLDER = "（空表）"
SCAN_HINT = "（本页未提取到文本，可能是扫描件，请用本地 MinerU）"

# 合成锚点（禁止真实业务正文）
ANCHOR_TXT = "V1J_TXT_BODY_SYNTH"
ANCHOR_MD = "V1J_MD_BODY_SYNTH"
ANCHOR_DOCX = "V1J_DOCX_BODY_SYNTH"
ANCHOR_PAGE1 = "V1J_PAGE1_SYNTH"
ANCHOR_PAGE2 = "V1J_PAGE2_SYNTH"
ANCHOR_EMPTY_TBL_BEFORE = "V1J_EMPTY_TBL_BEFORE"
ANCHOR_EMPTY_TBL_AFTER = "V1J_EMPTY_TBL_AFTER"
ANCHOR_KB = "V1J_KB_READY_ANCHOR_等保"
BASELINE_MD = "V1J-PROJECT-BASELINE-KEEP"

# 错误面禁止泄漏的片段
_FORBIDDEN_IN_ERROR = (
    "NoUsableTextError",
    "ValueError",
    "Error",
    "Traceback",
    "C:\\",
    "C:/",
    "/Users/",
    "\\Users\\",
    "biaoshu.db",
    "Cookie",
    "CSRF",
    "password",
    "api_key",
)

_TEMP_ROOTS: list[Path] = []


def _track_temp_root(root: Path) -> Path:
    """用途：登记 TEMP 根，清理后断言不存在。"""
    resolved = root.resolve()
    _TEMP_ROOTS.append(resolved)
    return resolved


def _make_temp_dir(prefix: str = "v1j_synth_") -> Path:
    return _track_temp_root(Path(tempfile.mkdtemp(prefix=prefix)))


def _cleanup_temp_root(root: Path) -> None:
    if root.exists():
        shutil.rmtree(root, ignore_errors=False)
    if root.exists():
        shutil.rmtree(root, ignore_errors=True)


def _assert_temp_gone(root: Path) -> None:
    assert not root.exists(), f"TEMP 根仍存在: {root}"


def _use_temp_upload_root(tmp_path: Path, monkeypatch) -> Path:
    """用途：upload_dir 指向 pytest TEMP，知识库根随之落到 TEMP 下 data/knowledge。"""
    root = tmp_path / "v1j_upload_root"
    root.mkdir(parents=True, exist_ok=True)
    settings = get_settings()
    monkeypatch.setattr(settings, "upload_dir", str(root.resolve()))
    return root.resolve()


def _pdf_escape_literal(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_two_page_text_pdf(page1: str, page2: str) -> bytes:
    """用途：标准库字节构造两页可提取文字 PDF（空串页视为 blank）。"""
    streams = []
    for text in (page1, page2):
        content = f"BT /F1 12 Tf 72 720 Td ({_pdf_escape_literal(text)}) Tj ET"
        streams.append(content.encode("latin-1"))

    objs: list[bytes] = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >>\nendobj\n",
        (
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 5 0 R /Resources << /Font << /F1 7 0 R >> >> >>\nendobj\n"
        ),
        (
            b"4 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 6 0 R /Resources << /Font << /F1 7 0 R >> >> >>\nendobj\n"
        ),
        (
            f"5 0 obj\n<< /Length {len(streams[0])} >>\nstream\n".encode("latin-1")
            + streams[0]
            + b"\nendstream\nendobj\n"
        ),
        (
            f"6 0 obj\n<< /Length {len(streams[1])} >>\nstream\n".encode("latin-1")
            + streams[1]
            + b"\nendstream\nendobj\n"
        ),
        b"7 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objs:
        offsets.append(len(out))
        out.extend(obj)
    xref_at = len(out)
    out.extend(f"xref\n0 {len(objs) + 1}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets:
        out.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        (
            f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n"
        ).encode("latin-1")
    )
    return bytes(out)


def build_zero_page_pdf() -> bytes:
    """用途：零页 PDF（Pages.Kids 空、Count=0）。"""
    objs: list[bytes] = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [] /Count 0 >>\nendobj\n",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for obj in objs:
        offsets.append(len(out))
        out.extend(obj)
    xref_at = len(out)
    out.extend(f"xref\n0 {len(objs) + 1}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets:
        out.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        (
            f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n"
        ).encode("latin-1")
    )
    return bytes(out)


def _write_empty_docx(path: Path) -> None:
    """用途：仅空段落，无可见正文。"""
    doc = Document()
    doc.add_paragraph("")
    doc.save(str(path))


def _write_empty_paragraphs_only_docx(path: Path) -> None:
    """用途：多个空段落。"""
    doc = Document()
    for _ in range(3):
        doc.add_paragraph("   ")
        doc.add_paragraph("\n")
    doc.save(str(path))


def _write_empty_table_only_docx(path: Path) -> None:
    """用途：仅空表（零行），无非空段落/单元格。"""
    doc = Document()
    # 默认空段忽略；唯一块为空表
    doc.add_table(rows=0, cols=2)
    doc.save(str(path))


def _write_body_plus_empty_table_docx(path: Path) -> None:
    """用途：正文 + 空表 + 正文，顺序与 V1-D 空表块一致。"""
    doc = Document()
    doc.add_paragraph(ANCHOR_EMPTY_TBL_BEFORE)
    doc.add_table(rows=0, cols=2)
    doc.add_paragraph(ANCHOR_EMPTY_TBL_AFTER)
    doc.save(str(path))


def _write_body_docx(path: Path, body: str = ANCHOR_DOCX) -> None:
    doc = Document()
    doc.add_paragraph(body)
    doc.save(str(path))


def _assert_no_usable_raise(path: Path, name: str) -> BaseException:
    """
    用途：锁固定中文 ValueError；异常出现后再断言精确类型为 NoUsableTextError。
    说明：禁止顶层 from-import NoUsableTextError（生产未实现时 collection error）。
    """
    with pytest.raises(ValueError, match=re.escape(FIXED_ERR)) as ei:
        parse_file_to_markdown(path, name)
    exc = ei.value
    assert str(exc) == FIXED_ERR, f"错误文案非精确相等: {str(exc)!r}"
    # 异常出现后才访问类；类型必须精确
    assert type(exc) is parse_service.NoUsableTextError, (
        f"异常类型非 NoUsableTextError: {type(exc)!r}"
    )
    blob = str(exc)
    for frag in _FORBIDDEN_IN_ERROR:
        assert frag not in blob, f"错误面泄漏片段 {frag!r}: {blob!r}"
    assert not re.search(r"[A-Za-z]:\\", blob), f"疑似绝对路径: {blob!r}"
    return exc


def _create_project(client: TestClient, name: str) -> str:
    res = client.post("/api/projects", json={"name": name, "kind": "technical"})
    assert res.status_code in (200, 201), res.text
    return res.json()["id"]


def _get_state(client: TestClient, pid: str) -> dict:
    res = client.get(f"/api/projects/{pid}/editor-state")
    assert res.status_code == 200, res.text
    return res.json()


def _put_state(client: TestClient, pid: str, body: dict) -> dict:
    res = client.put(f"/api/projects/{pid}/editor-state", json=body)
    assert res.status_code == 200, res.text
    return res.json()


def _get_project(client: TestClient, pid: str) -> dict:
    res = client.get(f"/api/projects/{pid}")
    assert res.status_code == 200, res.text
    return res.json()


def _full_editor_fingerprint(state: dict) -> dict[str, Any]:
    """
    用途：对 GET editor-state 完整 JSON 做独立深拷贝，供失败路径前后整包精确比较。
    覆盖 projectId、全部 13 个权威键、stateVersion、responseMatrixVersion、updatedAt、
    currentRevisionSourceKind/currentRevisionActorUsername 等响应稳定字段；禁止只选子集。
    """
    return copy.deepcopy(state)


def _project_step_fingerprint(proj: dict) -> dict[str, Any]:
    return {
        "status": proj.get("status"),
        "technicalPlanStep": proj.get("technicalPlanStep"),
    }


def _physical_chunk_count(doc_id: str) -> int:
    """用途：物理统计 KbChunkRow，禁止只看 chunk_count 假绿。"""
    db = SessionLocal()
    try:
        n = db.scalar(
            select(func.count())
            .select_from(KbChunkRow)
            .where(KbChunkRow.document_id == doc_id)
        )
        return int(n or 0)
    finally:
        db.close()


def _assert_status_message_clean(msg: str | None) -> None:
    assert msg == FIXED_ERR, f"statusMessage 非固定中文: {msg!r}"
    for frag in ("NoUsableTextError", "ValueError", "Traceback"):
        assert frag not in (msg or ""), f"statusMessage 含 {frag}: {msg!r}"
    assert not re.search(r"[A-Za-z]:\\", msg or ""), f"statusMessage 含路径: {msg!r}"
    assert not re.search(r"(?:/Users/|\\\\Users\\\\)", msg or ""), (
        f"statusMessage 含用户目录: {msg!r}"
    )


# ---------------------------------------------------------------------------
# 1) 运行时类存在性（生产未实现 → 业务红，非 import 红）
# ---------------------------------------------------------------------------
def test_no_usable_text_error_class_exists_runtime():
    """运行时 NoUsableTextError 必须存在且继承 ValueError。"""
    assert hasattr(parse_service, "NoUsableTextError"), (
        "parse_service 尚无 NoUsableTextError（生产未实现）"
    )
    cls = parse_service.NoUsableTextError
    assert isinstance(cls, type)
    assert issubclass(cls, ValueError)
    # 实例化后 str 可由调用方传入固定中文
    inst = cls(FIXED_ERR)
    assert isinstance(inst, ValueError)
    assert str(inst) == FIXED_ERR


# ---------------------------------------------------------------------------
# 2) 空/空白 TXT/MD/markdown
# ---------------------------------------------------------------------------
def test_empty_and_whitespace_txt_md_markdown_raise():
    """空字节、空白、CR/LF/TAB 的 .txt/.md/.markdown 均抛精确 NoUsableTextError。"""
    root = _make_temp_dir()
    try:
        payloads: list[bytes] = [
            b"",
            b" ",
            b"   ",
            b"\n",
            b"\r\n",
            b"\t",
            b" \n\t \r\n ",
            "\u3000\n".encode("utf-8"),  # 全角空格+换行
        ]
        suffixes = (".txt", ".md", ".markdown")
        for suffix in suffixes:
            for i, raw in enumerate(payloads):
                name = f"empty_{i}{suffix}"
                path = root / name
                path.write_bytes(raw)
                _assert_no_usable_raise(path, name)
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 3) 空 DOCX / 仅空段 / 仅空表
# ---------------------------------------------------------------------------
def test_empty_docx_variants_raise():
    """空 DOCX、仅空段落、仅空表均抛精确错误。"""
    root = _make_temp_dir()
    try:
        cases = (
            ("v1j_empty.docx", _write_empty_docx),
            ("v1j_empty_paras.docx", _write_empty_paragraphs_only_docx),
            ("v1j_empty_table_only.docx", _write_empty_table_only_docx),
        )
        for name, writer in cases:
            path = root / name
            writer(path)
            assert path.is_file() and path.stat().st_size > 0
            _assert_no_usable_raise(path, name)
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 4) 两页 blank PDF + 零页 PDF
# ---------------------------------------------------------------------------
def test_blank_and_zero_page_pdf_raise():
    """两页 blank 与零页 PDF 均抛精确 NoUsableTextError。"""
    root = _make_temp_dir()
    try:
        blank_name = "v1j_two_blank.pdf"
        blank_path = root / blank_name
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_blank_page(width=612, height=792)
        with blank_path.open("wb") as fh:
            writer.write(fh)
        _assert_no_usable_raise(blank_path, blank_name)

        zero_name = "v1j_zero_page.pdf"
        zero_path = root / zero_name
        zero_path.write_bytes(build_zero_page_pdf())
        _assert_no_usable_raise(zero_path, zero_name)
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 5) 有正文各格式继续 success
# ---------------------------------------------------------------------------
def test_nonempty_formats_still_success():
    """有正文 TXT/MD/DOCX/PDF 继续 success；文件名仅进标题，不误判固定错误。"""
    root = _make_temp_dir()
    try:
        txt_name = "v1j_body.txt"
        txt_path = root / txt_name
        txt_path.write_text(f"{ANCHOR_TXT}\n", encoding="utf-8")
        txt_md = parse_file_to_markdown(txt_path, txt_name)
        assert txt_md == f"{HEADER_PREFIX}{txt_name}\n\n{ANCHOR_TXT}\n"
        assert FIXED_ERR not in txt_md

        md_name = "v1j_body.md"
        md_path = root / md_name
        md_path.write_text(f"# Title\n\n{ANCHOR_MD}\n", encoding="utf-8")
        out_md = parse_file_to_markdown(md_path, md_name)
        assert out_md.startswith(f"{HEADER_PREFIX}{md_name}\n\n")
        assert ANCHOR_MD in out_md
        assert FIXED_ERR not in out_md

        docx_name = "v1j_body.docx"
        docx_path = root / docx_name
        _write_body_docx(docx_path)
        docx_md = parse_file_to_markdown(docx_path, docx_name)
        assert docx_md.startswith(f"{HEADER_PREFIX}{docx_name}\n\n")
        assert ANCHOR_DOCX in docx_md
        assert FIXED_ERR not in docx_md

        pdf_name = "v1j_body.pdf"
        pdf_path = root / pdf_name
        pdf_path.write_bytes(build_two_page_text_pdf(ANCHOR_PAGE1, ANCHOR_PAGE2))
        pdf_md = parse_file_to_markdown(pdf_path, pdf_name)
        assert pdf_md.startswith(f"{HEADER_PREFIX}{pdf_name}\n\n")
        assert ANCHOR_PAGE1 in pdf_md and ANCHOR_PAGE2 in pdf_md
        assert SCAN_HINT not in pdf_md
        assert FIXED_ERR not in pdf_md
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 6) 正文 + 空表继续 success
# ---------------------------------------------------------------------------
def test_docx_body_plus_empty_table_success_order():
    """正文加空表继续 success，保留正文与「（空表）」顺序。"""
    root = _make_temp_dir()
    try:
        name = "v1j_body_empty_table.docx"
        path = root / name
        _write_body_plus_empty_table_docx(path)
        md = parse_file_to_markdown(path, name)
        expected = (
            f"{HEADER_PREFIX}{name}\n\n"
            f"{ANCHOR_EMPTY_TBL_BEFORE}\n\n"
            f"{EMPTY_TABLE_PLACEHOLDER}\n\n"
            f"{ANCHOR_EMPTY_TBL_AFTER}\n"
        )
        assert md == expected
        assert FIXED_ERR not in md
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 7) 混合 PDF：真实锚点 + 精确空页 SCAN_HINT
# ---------------------------------------------------------------------------
def test_mixed_pdf_real_anchor_and_exact_blank_hint():
    """混合 PDF 至少一个真实文字锚点时 success；空页 SCAN_HINT 数量与顺序精确。"""
    root = _make_temp_dir()
    try:
        name = "v1j_mixed.pdf"
        path = root / name
        # 第 1 页有锚点，第 2 页空串 → blank
        path.write_bytes(build_two_page_text_pdf(ANCHOR_PAGE1, ""))
        md = parse_file_to_markdown(path, name)
        assert md.startswith(f"{HEADER_PREFIX}{name}\n\n")
        assert ANCHOR_PAGE1 in md
        assert FIXED_ERR not in md
        p1 = "## 第 1 页"
        p2 = "## 第 2 页"
        assert p1 in md and p2 in md
        assert md.count(SCAN_HINT) == 1
        i_p1 = md.index(p1)
        i_a1 = md.index(ANCHOR_PAGE1)
        i_p2 = md.index(p2)
        i_hint = md.index(SCAN_HINT)
        assert i_p1 < i_a1 < i_p2 < i_hint
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)


# ---------------------------------------------------------------------------
# 8) 项目 parse：代表空文 failed + 完整指纹不变
# ---------------------------------------------------------------------------
def test_project_parse_empty_representatives_failed_no_side_effects(
    client: TestClient, tmp_path: Path, monkeypatch
):
    """
    每类代表空文 parse 必须 failed：固定中文 error、result 空；
    editor-state 完整指纹与 project status/step 不变；零 analyze/outline/LLM 调用。
    """
    upload_root = _use_temp_upload_root(tmp_path, monkeypatch)
    assert upload_root.is_absolute()
    assert str(tmp_path.resolve()) in str(upload_root)

    analyze_calls: list[Any] = []
    outline_calls: list[Any] = []
    llm_calls: list[Any] = []

    def _forbid_analyze(*_a, **_k):
        analyze_calls.append(1)
        raise AssertionError("空文 parse 不得调用 _run_analyze")

    def _forbid_outline(*_a, **_k):
        outline_calls.append(1)
        raise AssertionError("空文 parse 不得调用 _run_outline")

    def _forbid_llm(*_a, **_k):
        llm_calls.append(1)
        raise AssertionError("空文 parse 不得调用 LLM")

    monkeypatch.setattr("app.services.task_service._run_analyze", _forbid_analyze)
    monkeypatch.setattr("app.services.task_service._run_outline", _forbid_outline)
    monkeypatch.setattr("app.services.llm_service.chat_completion", _forbid_llm)

    # 合成空文件字节
    empty_txt = b""
    empty_md = b"\n\t  \n"
    # 空 DOCX
    _tmp_docx = tmp_path / "_build_empty.docx"
    _write_empty_docx(_tmp_docx)
    empty_docx = _tmp_docx.read_bytes()
    # 两页 blank PDF
    pdf_buf = BytesIO()
    w = PdfWriter()
    w.add_blank_page(width=612, height=792)
    w.add_blank_page(width=612, height=792)
    w.write(pdf_buf)
    empty_pdf = pdf_buf.getvalue()

    samples: list[tuple[str, bytes, str]] = [
        ("empty.txt", empty_txt, "text/plain"),
        ("empty.md", empty_md, "text/markdown"),
        ("empty.docx", empty_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ("empty.pdf", empty_pdf, "application/pdf"),
    ]

    for filename, content, mime in samples:
        pid = _create_project(client, f"V1J-parse-{filename}")
        _put_state(
            client,
            pid,
            {
                "parsedMarkdown": BASELINE_MD,
                "analysisOverview": "V1J-overview-keep",
                "outline": [{"id": "n1", "title": "保留章", "children": []}],
                "chapters": [
                    {
                        "id": "n1",
                        "title": "保留章",
                        "body": "保留正文",
                        "preview": "保留正文",
                        "wordCount": 4,
                        "status": "done",
                        "targetWords": 100,
                    }
                ],
            },
        )

        up = client.post(
            f"/api/projects/{pid}/files",
            files={"file": (filename, BytesIO(content), mime)},
        )
        assert up.status_code == 201, up.text

        # 上传后 fingerprint 以当前 GET editor-state/project 整包为准（上传不改 parsedMarkdown）
        before_state = _full_editor_fingerprint(_get_state(client, pid))
        before_proj = _project_step_fingerprint(_get_project(client, pid))

        parse_task = client.post(
            f"/api/projects/{pid}/tasks?sync=true",
            json={"type": "parse"},
        )
        assert parse_task.status_code == 201, parse_task.text
        body = parse_task.json()
        assert body["status"] == "failed", body
        assert body["message"] == "任务失败", body
        assert body["error"] == FIXED_ERR, body
        assert body.get("result") in (None, {}, []), body
        for frag in ("NoUsableTextError", "ValueError", "Traceback"):
            assert frag not in json.dumps(body, ensure_ascii=False)

        after_state = _full_editor_fingerprint(_get_state(client, pid))
        after_proj = _project_step_fingerprint(_get_project(client, pid))
        assert after_state == before_state, (
            f"{filename}: editor-state 被改写 before={before_state} after={after_state}"
        )
        assert after_proj == before_proj, (
            f"{filename}: project status/step 被改写 before={before_proj} after={after_proj}"
        )
        assert after_state["parsedMarkdown"] == BASELINE_MD

    assert analyze_calls == [], f"发生 analyze 调用: {analyze_calls}"
    assert outline_calls == [], f"发生 outline 调用: {outline_calls}"
    assert llm_calls == [], f"发生 LLM 调用: {llm_calls}"

    # 样本与 upload 均在 TEMP
    assert upload_root.exists()
    # 不在此删 pytest tmp_path（由 pytest 回收）；显式确认非仓库 uploads
    assert "v1j_upload_root" in str(upload_root)


# ---------------------------------------------------------------------------
# 9) 知识库：新空文 failed；ready→空文重索引物理清 chunk；有正文仍 ready
# ---------------------------------------------------------------------------
def test_knowledge_empty_fail_and_reindex_clears_physical_chunks(
    client: TestClient, tmp_path: Path, monkeypatch
):
    """
    新空文：保留 doc/file、failed、固定纯中文、零 chunk；
    ready 文档把 TEMP 中 stored file 改空后重索引，旧 KbChunkRow 物理为 0；
    有正文仍 ready 可检索。
    """
    upload_root = _use_temp_upload_root(tmp_path, monkeypatch)

    # --- 新空文 ---
    empty_up = client.post(
        "/api/knowledge/docs/upload",
        files={"file": ("v1j-empty.txt", BytesIO(b""), "text/plain")},
    )
    assert empty_up.status_code == 201, empty_up.text
    empty_doc = empty_up.json()
    empty_id = empty_doc["id"]
    assert empty_doc["status"] == "failed", empty_doc
    assert empty_doc.get("statusMessage") == FIXED_ERR, empty_doc
    assert empty_doc.get("chunks") == 0, empty_doc
    _assert_status_message_clean(empty_doc.get("statusMessage"))
    assert _physical_chunk_count(empty_id) == 0

    # 文档行与磁盘文件仍在
    db = SessionLocal()
    try:
        row = db.get(KbDocumentRow, empty_id)
        assert row is not None
        assert row.status == "failed"
        assert row.status_message == FIXED_ERR
        assert row.chunk_count == 0
        assert row.stored_name
        settings = get_settings()
        from app.services import knowledge_service as ks

        disk = ks._doc_dir(settings, row.workspace_id, row.id) / row.stored_name
        assert disk.is_file(), f"空文失败后文件应保留: {disk}"
        assert str(upload_root.parent) in str(disk.resolve()) or str(tmp_path.resolve()) in str(
            disk.resolve()
        )
    finally:
        db.close()

    # --- ready 有正文 → 改空 → 重索引 ---
    ready_content = f"# KB\n\n{ANCHOR_KB} 合成可检索正文。\n".encode("utf-8")
    ready_up = client.post(
        "/api/knowledge/docs/upload",
        files={"file": ("v1j-ready.md", BytesIO(ready_content), "text/markdown")},
    )
    assert ready_up.status_code == 201, ready_up.text
    ready_doc = ready_up.json()
    ready_id = ready_doc["id"]
    assert ready_doc["status"] == "ready", ready_doc
    assert ready_doc["chunks"] >= 1
    assert _physical_chunk_count(ready_id) >= 1

    search_ok = client.get("/api/knowledge/search", params={"q": "等保", "topK": 5})
    assert search_ok.status_code == 200
    assert search_ok.json()["count"] >= 1

    # 覆盖 TEMP 中该文档自己的 stored file 为空
    db = SessionLocal()
    try:
        row = db.get(KbDocumentRow, ready_id)
        assert row is not None and row.stored_name
        settings = get_settings()
        from app.services import knowledge_service as ks

        disk = ks._doc_dir(settings, row.workspace_id, row.id) / row.stored_name
        assert disk.is_file()
        disk.write_bytes(b"")
        assert disk.stat().st_size == 0
    finally:
        db.close()

    ri = client.post(f"/api/knowledge/docs/{ready_id}/reindex")
    assert ri.status_code == 200, ri.text
    re_doc = ri.json()
    assert re_doc["status"] == "failed", re_doc
    assert re_doc.get("statusMessage") == FIXED_ERR, re_doc
    assert re_doc.get("chunks") == 0, re_doc
    _assert_status_message_clean(re_doc.get("statusMessage"))
    # 物理块必须为 0（禁止只改计数）
    assert _physical_chunk_count(ready_id) == 0

    # 文件仍在
    db = SessionLocal()
    try:
        row = db.get(KbDocumentRow, ready_id)
        assert row is not None
        settings = get_settings()
        from app.services import knowledge_service as ks

        disk = ks._doc_dir(settings, row.workspace_id, row.id) / (row.stored_name or "")
        assert disk.is_file()
    finally:
        db.close()

    # --- 另传有正文，仍 ready 可检索 ---
    body2 = f"# OK\n\n{ANCHOR_KB} 第二篇可检索。\n".encode("utf-8")
    ok_up = client.post(
        "/api/knowledge/docs/upload",
        files={"file": ("v1j-ok2.md", BytesIO(body2), "text/markdown")},
    )
    assert ok_up.status_code == 201, ok_up.text
    ok_doc = ok_up.json()
    assert ok_doc["status"] == "ready"
    assert ok_doc.get("statusMessage") in (None, "")
    assert ok_doc["chunks"] >= 1
    assert _physical_chunk_count(ok_doc["id"]) >= 1
    search2 = client.get("/api/knowledge/search", params={"q": "等保", "topK": 5})
    assert search2.status_code == 200
    assert search2.json()["count"] >= 1

    assert upload_root.exists()


# ---------------------------------------------------------------------------
# 10) TEMP 登记与清理（本模块自建 TEMP）
# ---------------------------------------------------------------------------
def test_v1j_temp_roots_tracked_and_cleaned():
    """TEMP 根显式登记且清理后不存在；禁止空列表假绿。"""
    before = len(_TEMP_ROOTS)
    root = _make_temp_dir()
    try:
        assert len(_TEMP_ROOTS) == before + 1
        assert _TEMP_ROOTS[-1] == root.resolve()
        assert root.is_absolute()
        assert "v1j_synth_" in root.name
        marker = root / "marker.txt"
        marker.write_text("V1J_TEMP\n", encoding="utf-8")
        assert marker.is_file()
    finally:
        _cleanup_temp_root(root)
        _assert_temp_gone(root)

    still = [str(p) for p in _TEMP_ROOTS if p.exists()]
    assert still == [], f"仍有未清理 TEMP: {still}"
