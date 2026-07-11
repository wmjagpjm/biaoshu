"""
模块：项目正文图片测试
用途：验收图片角色隔离、受控 Markdown 引用、Word 嵌图、无效引用降级和项目删除清理。
对接：/api/projects/{id}/images、/files、任务 parse/export、file_service、export_service。
二次开发：新增图片格式、远程资源或严格失败策略时，必须补越权、路径和导出回归用例。
"""

from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path

from PIL import Image
import pytest

from app.core.config import get_settings
from app.core.database import SessionLocal
from app.models.entities import Project
from app.services import file_service, project_service


def _png_bytes() -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (8, 4), (34, 102, 170)).save(buffer, format="PNG")
    return buffer.getvalue()


def _image_bytes(image_format: str) -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (8, 4), (34, 102, 170)).save(buffer, format=image_format)
    return buffer.getvalue()


def _create_project_with_source(client, name: str) -> str:
    project = client.post("/api/projects", json={"name": name}).json()
    project_id = project["id"]
    source = client.post(
        f"/api/projects/{project_id}/files",
        files={
            "file": (
                "tender.md",
                BytesIO("# 招标文件\n\n智慧园区项目。".encode("utf-8")),
                "text/markdown",
            )
        },
    )
    assert source.status_code == 201
    return project_id


def test_image_upload_isolated_from_source_file_and_parse(client):
    project_id = _create_project_with_source(client, "图片角色隔离")

    image = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("topology.png", BytesIO(_png_bytes()), "image/png")},
    )
    assert image.status_code == 201
    image_id = image.json()["id"]

    source_files = client.get(f"/api/projects/{project_id}/files")
    assert source_files.status_code == 200
    assert [item["filename"] for item in source_files.json()] == ["tender.md"]

    images = client.get(f"/api/projects/{project_id}/images")
    assert images.status_code == 200
    assert [item["id"] for item in images.json()] == [image_id]

    parsed = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "parse"},
    )
    assert parsed.status_code == 201
    assert parsed.json()["status"] == "success"
    state = client.get(f"/api/projects/{project_id}/editor-state").json()
    assert "智慧园区项目" in state["parsedMarkdown"]


def test_image_upload_rejects_fake_image(client):
    project_id = _create_project_with_source(client, "伪装图片")

    response = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("not-image.png", BytesIO(b"not an image"), "image/png")},
    )

    assert response.status_code == 400
    assert "图片" in response.json()["detail"]


@pytest.mark.parametrize(
    ("image_format", "filename", "content_type"),
    [
        ("JPEG", "topology.jpg", "image/jpeg"),
        ("GIF", "topology.gif", "image/gif"),
    ],
)
def test_image_upload_accepts_supported_non_png_formats(
    client, image_format, filename, content_type
):
    project_id = _create_project_with_source(client, f"{image_format} 图片")

    response = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": (filename, BytesIO(_image_bytes(image_format)), content_type)},
    )

    assert response.status_code == 201
    assert response.json()["contentType"] == content_type


def test_image_upload_rejects_unsupported_format_and_limits(client, monkeypatch):
    project_id = _create_project_with_source(client, "图片边界")
    unsupported = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("topology.bmp", BytesIO(_image_bytes("BMP")), "image/bmp")},
    )
    assert unsupported.status_code == 400
    assert "仅支持" in unsupported.json()["detail"]

    settings = get_settings()
    image_bytes = _png_bytes()
    monkeypatch.setattr(settings, "max_image_upload_bytes", len(image_bytes) - 1)
    byte_limited = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("limited.png", BytesIO(image_bytes), "image/png")},
    )
    assert byte_limited.status_code == 400
    assert "图片过大" in byte_limited.json()["detail"]

    monkeypatch.setattr(settings, "max_image_upload_bytes", len(image_bytes) + 1)
    monkeypatch.setattr(settings, "max_image_pixels", 31)
    pixel_limited = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("pixel.png", BytesIO(image_bytes), "image/png")},
    )
    assert pixel_limited.status_code == 400
    assert "像素" in pixel_limited.json()["detail"]


def test_image_limit_is_serialized_across_concurrent_sessions(client, monkeypatch):
    project_id = _create_project_with_source(client, "图片并发限额")
    settings = get_settings()
    monkeypatch.setattr(settings, "max_project_images", 1)

    def upload(index: int) -> tuple[str, str]:
        db = SessionLocal()
        try:
            row = file_service.save_image_upload(
                db,
                "ws_local",
                project_id,
                settings,
                filename=f"concurrent-{index}.png",
                content=_png_bytes(),
            )
            return "success", row.id
        except ValueError as exc:
            return "rejected", str(exc)
        finally:
            db.close()

    with ThreadPoolExecutor(max_workers=2) as executor:
        results = list(executor.map(upload, range(2)))

    assert [status for status, _ in results].count("success") == 1
    assert [status for status, _ in results].count("rejected") == 1
    assert len(client.get(f"/api/projects/{project_id}/images").json()) == 1


def test_image_upload_accepts_session_with_prior_read(client):
    project_id = _create_project_with_source(client, "图片事务复用")
    db = SessionLocal()
    try:
        assert db.get(Project, project_id) is not None
        row = file_service.save_image_upload(
            db,
            "ws_local",
            project_id,
            get_settings(),
            filename="active-session.png",
            content=_png_bytes(),
        )
    finally:
        db.close()

    assert row.role == file_service.FILE_ROLE_IMAGE


def test_failed_image_commit_removes_orphan_file(client, monkeypatch):
    project_id = _create_project_with_source(client, "图片提交清理")
    project_dir = Path(get_settings().upload_dir) / project_id
    before = {path.name for path in project_dir.iterdir()}
    db = SessionLocal()

    def fail_commit():
        raise RuntimeError("模拟提交失败")

    monkeypatch.setattr(db, "commit", fail_commit)
    try:
        with pytest.raises(RuntimeError, match="模拟提交失败"):
            file_service.save_image_upload(
                db,
                "ws_local",
                project_id,
                get_settings(),
                filename="orphan.png",
                content=_png_bytes(),
            )
    finally:
        db.close()

    assert {path.name for path in project_dir.iterdir()} == before


def test_failed_image_commit_logs_orphan_cleanup_failure(client, monkeypatch, caplog):
    project_id = _create_project_with_source(client, "图片孤儿告警")
    project_dir = Path(get_settings().upload_dir) / project_id
    before = {path.name for path in project_dir.iterdir()}
    real_unlink = Path.unlink
    db = SessionLocal()

    def fail_commit():
        raise RuntimeError("模拟提交失败")

    def fail_unlink(*_args, **_kwargs):
        raise OSError("模拟文件占用")

    monkeypatch.setattr(db, "commit", fail_commit)
    monkeypatch.setattr(file_service.Path, "unlink", fail_unlink)
    try:
        with caplog.at_level("WARNING", logger=file_service.__name__):
            with pytest.raises(RuntimeError, match="模拟提交失败"):
                file_service.save_image_upload(
                    db,
                    "ws_local",
                    project_id,
                    get_settings(),
                    filename="orphan-log.png",
                    content=_png_bytes(),
                )
    finally:
        db.close()

    orphan_paths = [path for path in project_dir.iterdir() if path.name not in before]
    assert len(orphan_paths) == 1
    assert "清理孤儿文件失败" in caplog.text
    real_unlink(orphan_paths[0], missing_ok=True)


def test_image_rejects_cross_project_and_source_file_references(client):
    image_project_id = _create_project_with_source(client, "图片归属项目")
    target_project_id = _create_project_with_source(client, "图片引用项目")
    image = client.post(
        f"/api/projects/{image_project_id}/images",
        files={"file": ("topology.png", BytesIO(_png_bytes()), "image/png")},
    )
    assert image.status_code == 201
    image_id = image.json()["id"]
    source_id = client.get(
        f"/api/projects/{target_project_id}/files"
    ).json()[0]["id"]

    assert client.get(
        f"/api/projects/{target_project_id}/images/{image_id}"
    ).status_code == 404
    assert client.get(
        f"/api/projects/{target_project_id}/images/{source_id}"
    ).status_code == 404

    client.put(
        f"/api/projects/{target_project_id}/editor-state",
        json={
            "chapters": [
                {
                    "id": "chapter-1",
                    "title": "图片隔离",
                    "body": (
                        f"![跨项目](biaoshu-image://{image_id})\n"
                        f"![源文件](biaoshu-image://{source_id})"
                    ),
                    "preview": "图片隔离",
                    "wordCount": 2,
                    "status": "done",
                }
            ]
        },
    )
    exported = client.post(
        f"/api/projects/{target_project_id}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    assert len(exported.json()["result"]["imageWarnings"]) == 2

    from docx import Document  # type: ignore

    downloaded = client.get(
        f"/api/projects/{target_project_id}/export/download/"
        f"{exported.json()['result']['storedName']}"
    )
    doc = Document(BytesIO(downloaded.content))
    assert len(doc.inline_shapes) == 0


def test_image_path_rejects_directory_escape():
    with pytest.raises(ValueError, match="非法存储文件名"):
        file_service.resolve_path(
            get_settings(),
            "proj_image_path_safety",
            "../outside.png",
        )


def test_export_embeds_project_image_and_warns_for_invalid_reference(client):
    project_id = _create_project_with_source(client, "图片导出")
    image = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("topology.png", BytesIO(_png_bytes()), "image/png")},
    )
    assert image.status_code == 201
    image_id = image.json()["id"]

    client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "image": {
                    "max_width_percent": 60,
                    "alignment": "居中对齐",
                    "caption_font": "黑体",
                    "caption_size": "五号",
                    "caption_alignment": "居中对齐",
                    "caption_bold": True,
                    "caption_italic": False,
                }
            }
        },
    )
    client.put(
        f"/api/projects/{project_id}/editor-state",
        json={
            "chapters": [
                {
                    "id": "chapter-1",
                    "title": "总体架构",
                    "body": (
                        f'![拓扑图](biaoshu-image://{image_id} "图 1 总体拓扑")\n'
                        "![非法图](biaoshu-image://../outside)"
                    ),
                    "preview": "图片",
                    "wordCount": 2,
                    "status": "done",
                }
            ]
        },
    )

    exported = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    result = exported.json()["result"]
    assert result["imageWarnings"]
    downloaded = client.get(
        f"/api/projects/{project_id}/export/download/{result['storedName']}"
    )
    assert downloaded.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    assert len(doc.inline_shapes) == 1
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "图 1 总体拓扑" in text
    assert "图片引用无效" in text


def test_delete_project_removes_project_upload_directory(client):
    project_id = _create_project_with_source(client, "图片目录清理")
    image = client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("topology.png", BytesIO(_png_bytes()), "image/png")},
    )
    assert image.status_code == 201
    project_dir = get_settings().upload_dir

    deleted = client.delete(f"/api/projects/{project_id}")

    assert deleted.status_code == 204
    assert not (Path(project_dir) / project_id).exists()


def test_delete_project_logs_upload_cleanup_failure(client, monkeypatch, caplog):
    project_id = _create_project_with_source(client, "图片清理告警")
    client.post(
        f"/api/projects/{project_id}/images",
        files={"file": ("topology.png", BytesIO(_png_bytes()), "image/png")},
    )
    project_dir = Path(get_settings().upload_dir) / project_id
    real_rmtree = project_service.shutil.rmtree

    def fail_rmtree(*_args, **_kwargs):
        raise OSError("模拟目录占用")

    monkeypatch.setattr(project_service.shutil, "rmtree", fail_rmtree)
    with caplog.at_level("WARNING", logger=project_service.__name__):
        deleted = client.delete(f"/api/projects/{project_id}")

    assert deleted.status_code == 204
    assert "清理上传目录失败" in caplog.text
    real_rmtree(project_dir, ignore_errors=True)
