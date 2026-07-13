"""
模块：Word 导出标题边框测试
用途：验收标题段落描边、分级底色、叶子标题左栏、关闭状态及 camelCase 配置兼容。
对接：export_service 标题渲染；设置 exportFormat；export 同步任务。
二次开发：structure / 整章页框不属于本测试范围；新增版式语义需另建用例。
"""

from io import BytesIO
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from app.services.export_service import (
    _apply_heading_border,
    _heading_border_cfg,
    _leaf_flags_from_levels,
    write_markdown_body,
)

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _left_border_attrs(paragraph) -> dict[str, str]:
    """用途：解析标题段落 w:pBdr/w:left 的 sz/space/color。"""
    root = ET.fromstring(paragraph._p.xml)
    left = root.find(".//w:pBdr/w:left", W_NS)
    if left is None:
        return {}
    return {
        "sz": left.get(f"{{{W_NS['w']}}}sz") or left.get("sz") or "",
        "space": left.get(f"{{{W_NS['w']}}}space") or left.get("space") or "",
        "color": left.get(f"{{{W_NS['w']}}}color") or left.get("color") or "",
    }


def _docx_word_xmls(content: bytes) -> dict[str, bytes]:
    """用途：读取 docx 压缩包内的 Word XML，核验页级 OOXML 属性。"""
    with ZipFile(BytesIO(content)) as archive:
        return {
            name: archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        }


def test_heading_border_config_accepts_camel_case_and_sanitizes_colors():
    cfg = _heading_border_cfg(
        {
            "headingBorder": {
                "enabled": True,
                "minHeadingLeftEnabled": True,
                "borderColor": "#123abc",
                "levelCellColors": ["#abcdef", "不是颜色"],
            }
        }
    )

    assert cfg["enabled"] is True
    assert cfg["min_heading_left_enabled"] is True
    assert cfg["border_color"] == "123ABC"
    assert cfg["level_cell_colors"] == ["ABCDEF", "FFFFFF"]

    disabled = _heading_border_cfg({"headingBorder": {"enabled": "false"}})
    enabled = _heading_border_cfg({"headingBorder": {"enabled": "true"}})
    assert disabled["enabled"] is False
    assert enabled["enabled"] is True

    snake = _heading_border_cfg(
        {
            "heading_border": {
                "enabled": True,
                "min_heading_left_enabled": "true",
            }
        }
    )
    assert snake["min_heading_left_enabled"] is True

    bogus = _heading_border_cfg(
        {
            "heading_border": {
                "enabled": True,
                "min_heading_left_enabled": "maybe",
            }
        }
    )
    assert bogus["min_heading_left_enabled"] is False


def test_leaf_flags_from_levels_detects_parents_and_leaves():
    # 一级 → 二级 → 三级：仅三级为叶；随后另一二级为叶
    assert _leaf_flags_from_levels([1, 2, 3, 2]) == [False, False, True, True]
    # 仅有一级
    assert _leaf_flags_from_levels([1]) == [True]
    # 四级存在时三级为父
    assert _leaf_flags_from_levels([2, 3, 4]) == [False, False, True]


def test_apply_heading_border_writes_paragraph_ooxml():
    from docx import Document  # type: ignore

    doc = Document()
    paragraph = doc.add_heading("一级标题", level=1)

    _apply_heading_border(
        paragraph,
        {
            "enabled": True,
            "border_color": "123456",
            "level_cell_colors": ["ABCDEF"],
        },
        level_index=0,
    )

    xml = paragraph._p.xml
    assert "w:pBdr" in xml
    assert 'w:color="123456"' in xml
    assert 'w:fill="ABCDEF"' in xml
    for edge_name in ("top", "left", "bottom", "right"):
        assert f"w:{edge_name}" in xml
    left = _left_border_attrs(paragraph)
    assert left.get("sz") == "8"
    assert left.get("space") == "4"


def test_apply_heading_border_leaf_left_emphasis():
    from docx import Document  # type: ignore

    doc = Document()
    leaf = doc.add_heading("叶子标题", level=3)
    parent = doc.add_heading("父标题", level=2)

    cfg = {
        "enabled": True,
        "min_heading_left_enabled": True,
        "border_color": "2468AC",
        "level_cell_colors": ["DDEEFF", "EEF6FF", "CCDDEE"],
    }
    _apply_heading_border(leaf, cfg, level_index=2, is_leaf=True)
    _apply_heading_border(parent, cfg, level_index=1, is_leaf=False)

    leaf_left = _left_border_attrs(leaf)
    parent_left = _left_border_attrs(parent)
    assert leaf_left.get("sz") == "18"
    assert leaf_left.get("space") == "6"
    assert leaf_left.get("color") == "2468AC"
    assert parent_left.get("sz") == "8"
    assert parent_left.get("space") == "4"
    # 同一 pBdr 内仅一条 left，不得叠双 left
    assert leaf._p.xml.count("<w:left ") == 1


def test_apply_heading_border_requires_both_switches_for_left_emphasis():
    from docx import Document  # type: ignore

    doc = Document()
    paragraph = doc.add_heading("仅开一边", level=2)

    _apply_heading_border(
        paragraph,
        {
            "enabled": True,
            "min_heading_left_enabled": False,
            "border_color": "123456",
            "level_cell_colors": ["ABCDEF"],
        },
        level_index=1,
        is_leaf=True,
    )
    left = _left_border_attrs(paragraph)
    assert left.get("sz") == "8"

    paragraph2 = doc.add_heading("边框关", level=2)
    _apply_heading_border(
        paragraph2,
        {
            "enabled": False,
            "min_heading_left_enabled": True,
            "border_color": "123456",
            "level_cell_colors": ["ABCDEF"],
        },
        level_index=1,
        is_leaf=True,
    )
    assert "w:pBdr" not in paragraph2._p.xml


def test_heading_border_keeps_ooxml_property_order():
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()
    paragraph = doc.add_heading("已有段落属性", level=1)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _apply_heading_border(
        paragraph,
        {
            "enabled": True,
            "border_color": "123456",
            "level_cell_colors": ["ABCDEF"],
        },
        level_index=0,
    )

    properties = paragraph._p.get_or_add_pPr()
    child_names = [child.tag.rsplit("}", 1)[-1] for child in properties]
    assert child_names.index("pStyle") < child_names.index("pBdr")
    assert child_names.index("pBdr") < child_names.index("shd")
    assert child_names.index("shd") < child_names.index("spacing")
    assert child_names.index("spacing") < child_names.index("jc")


def test_apply_heading_border_disabled_keeps_paragraph_plain():
    from docx import Document  # type: ignore

    doc = Document()
    paragraph = doc.add_heading("普通标题", level=1)

    _apply_heading_border(
        paragraph,
        {
            "enabled": False,
            "border_color": "123456",
            "level_cell_colors": ["ABCDEF"],
        },
        level_index=0,
    )

    xml = paragraph._p.xml
    assert "w:pBdr" not in xml
    assert "w:shd" not in xml


def test_markdown_heading_uses_matching_level_background():
    from docx import Document  # type: ignore

    doc = Document()
    write_markdown_body(
        doc,
        "## 部署拓扑\n正文内容",
        heading_border_cfg={
            "enabled": True,
            "border_color": "123456",
            "level_cell_colors": ["DDEEFF", "EEF6FF", "CCDDEE"],
        },
    )

    heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "部署拓扑")
    assert 'w:fill="CCDDEE"' in heading._p.xml


def test_markdown_leaf_headings_get_left_emphasis_only():
    from docx import Document  # type: ignore

    doc = Document()
    write_markdown_body(
        doc,
        "# 一级\n## 二级\n### 三级末级\n正文\n## 另一二级叶\n内容",
        heading_border_cfg={
            "enabled": True,
            "min_heading_left_enabled": True,
            "border_color": "112233",
            "level_cell_colors": ["AAAAAA", "BBBBBB", "CCCCCC", "DDDDDD"],
        },
    )

    by_text = {paragraph.text: paragraph for paragraph in doc.paragraphs}
    # Markdown 映射：#→H2, ##→H3, ###→H4
    assert _left_border_attrs(by_text["一级"]).get("sz") == "8"
    assert _left_border_attrs(by_text["二级"]).get("sz") == "8"
    assert _left_border_attrs(by_text["三级末级"]).get("sz") == "18"
    assert _left_border_attrs(by_text["三级末级"]).get("space") == "6"
    assert _left_border_attrs(by_text["另一二级叶"]).get("sz") == "18"


def test_export_heading_border_applies_to_generated_heading(client):
    settings = client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "template_name": "标题边框测试",
                "heading_border": {
                    "enabled": True,
                    "border_color": "#2468AC",
                    "level_cell_colors": [
                        "#DDEEFF",
                        "#EEF6FF",
                        "#FFFFFF",
                        "#FFFFFF",
                        "#FFFFFF",
                        "#FFFFFF",
                    ],
                },
            }
        },
    )
    assert settings.status_code == 200

    project = client.post("/api/projects", json={"name": "标题边框导出"}).json()
    project_id = project["id"]
    state = client.put(
        f"/api/projects/{project_id}/editor-state",
        json={
            "analysisOverview": "用于检查一级标题边框。",
            "chapters": [
                {
                    "id": "chapter-1",
                    "title": "实施方案",
                    "body": "## 部署拓扑\n正文内容",
                    "preview": "正文内容",
                    "wordCount": 4,
                    "status": "done",
                }
            ],
        },
    )
    assert state.status_code == 200

    exported = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(
        f"/api/projects/{project_id}/export/download/{stored_name}"
    )
    assert downloaded.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    overview_heading = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text == "一、项目概述 / 招标分析"
    )
    xml = overview_heading._p.xml
    assert "w:pBdr" in xml
    assert 'w:color="2468AC"' in xml
    assert 'w:fill="DDEEFF"' in xml
    assert all(
        b"pgBorders" not in xml
        for xml in _docx_word_xmls(downloaded.content).values()
    )


def test_export_min_heading_left_on_leaf_markdown(client):
    client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "template_name": "左栏测试",
                "heading_border": {
                    "enabled": True,
                    "min_heading_left_enabled": True,
                    "border_color": "#2468AC",
                    "level_cell_colors": ["#DDEEFF", "#EEF6FF", "#CCDDEE"],
                },
            }
        },
    )
    project = client.post("/api/projects", json={"name": "左栏导出"}).json()
    project_id = project["id"]
    client.put(
        f"/api/projects/{project_id}/editor-state",
        json={
            "analysisOverview": "概述。",
            "chapters": [
                {
                    "id": "chapter-1",
                    "title": "实施方案",
                    "body": "## 部署拓扑\n### 机房节点\n正文",
                    "preview": "正文",
                    "wordCount": 2,
                    "status": "done",
                }
            ],
        },
    )

    exported = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(
        f"/api/projects/{project_id}/export/download/{stored_name}"
    )
    assert downloaded.status_code == 200
    # 页面级边框不得出现
    assert all(
        b"pgBorders" not in xml
        for xml in _docx_word_xmls(downloaded.content).values()
    )

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    by_text = {paragraph.text: paragraph for paragraph in doc.paragraphs}
    # 固定概述/正文容器与章标题均不属于最小标题；Markdown ## 下有 ### → 非叶。
    assert _left_border_attrs(
        by_text["一、项目概述 / 招标分析"]
    ).get("sz") == "8"
    chapter = next(p for t, p in by_text.items() if "实施方案" in t)
    assert _left_border_attrs(chapter).get("sz") == "8"
    assert _left_border_attrs(by_text["部署拓扑"]).get("sz") == "8"
    assert _left_border_attrs(by_text["机房节点"]).get("sz") == "18"
    assert _left_border_attrs(by_text["机房节点"]).get("space") == "6"


def test_export_outline_leaf_gets_left_emphasis_only(client):
    client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "template_name": "大纲左栏测试",
                "heading_border": {
                    "enabled": True,
                    "min_heading_left_enabled": True,
                    "border_color": "#2468AC",
                    "level_cell_colors": ["#DDEEFF", "#EEF6FF", "#CCDDEE"],
                },
            }
        },
    )
    project = client.post("/api/projects", json={"name": "大纲左栏导出"}).json()
    project_id = project["id"]
    client.put(
        f"/api/projects/{project_id}/editor-state",
        json={
            "outline": [
                {
                    "id": "outline-parent",
                    "title": "总体设计",
                    "children": [
                        {"id": "outline-leaf", "title": "部署架构", "children": []}
                    ],
                }
            ]
        },
    )
    exported = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(
        f"/api/projects/{project_id}/export/download/{stored_name}"
    )
    assert downloaded.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    parent = next(p for p in doc.paragraphs if "总体设计" in p.text)
    leaf = next(p for p in doc.paragraphs if "部署架构" in p.text)
    assert _left_border_attrs(parent).get("sz") == "8"
    assert _left_border_attrs(leaf).get("sz") == "18"
    assert _left_border_attrs(leaf).get("space") == "6"
    assert all(
        b"pgBorders" not in xml
        for xml in _docx_word_xmls(downloaded.content).values()
    )


def test_business_export_heading_border_is_applied(client):
    client.put(
        "/api/settings",
        json={
            "exportFormat": {
                "template_name": "商务标题边框测试",
                "heading_border": {
                    "enabled": True,
                    "min_heading_left_enabled": True,
                    "border_color": "#3579BD",
                    "level_cell_colors": ["#DDEEFF", "#E6F2FF"],
                },
            }
        },
    )
    project = client.post(
        "/api/projects",
        json={"name": "商务边框导出", "kind": "business"},
    ).json()
    project_id = project["id"]
    client.put(
        f"/api/projects/{project_id}/editor-state",
        json={
            "businessQualify": [
                {
                    "id": "qualify-1",
                    "requirement": "法人资格",
                    "response": "已响应",
                    "evidence": "营业执照",
                    "status": "matched",
                }
            ]
        },
    )

    exported = client.post(
        f"/api/projects/{project_id}/tasks?sync=true",
        json={"type": "export", "payload": {"mode": "business"}},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(
        f"/api/projects/{project_id}/export/download/{stored_name}"
    )
    assert downloaded.status_code == 200
    assert all(
        b"pgBorders" not in xml
        for xml in _docx_word_xmls(downloaded.content).values()
    )

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    by_text = {paragraph.text: paragraph for paragraph in doc.paragraphs}
    # 文档名标题为父节点（下有 ## 资格响应），不得左栏强化
    assert "w:pBdr" in by_text["商务边框导出"]._p.xml
    assert _left_border_attrs(by_text["商务边框导出"]).get("sz") == "8"
    # 叶子小节「二、资格响应」应强化左栏
    assert _left_border_attrs(by_text["二、资格响应"]).get("sz") == "18"
    assert _left_border_attrs(by_text["二、资格响应"]).get("space") == "6"
