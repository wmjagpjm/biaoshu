"""
模块：响应矩阵 editor-state 与智能建议测试
用途：验收评分点/技术要求映射读写、容错、旧库补列，以及 response_match 候选分批与只读建议。
对接：GET|PUT /api/projects/{id}/editor-state；POST .../tasks type=response_match。
二次开发：新增矩阵字段或分批语义时必须补兼容/越界/不写库用例，避免防抖 PUT 清空已有映射。
"""

import json
from io import BytesIO

from sqlalchemy import create_engine, text

from app.core.database import SessionLocal, ensure_schema_columns
from app.models.entities import ProjectEditorStateRow
from app.services.llm_service import ChatResult


def _create_project(client) -> str:
    """用途：创建技术标项目并返回项目 id。"""
    return client.post("/api/projects", json={"name": "响应矩阵项目"}).json()["id"]


def test_response_matrix_empty_and_roundtrip(client):
    pid = _create_project(client)

    empty = client.get(f"/api/projects/{pid}/editor-state")
    assert empty.status_code == 200
    assert empty.json()["responseMatrix"] == []

    matrix = [
        {
            "id": "mx_req_1",
            "kind": "requirement",
            "sourceKey": "requirement:等保三级",
            "sourceIndex": 0,
            "sourceText": "等保三级",
            "weight": "",
            "chapterIds": ["chap_1"],
            "outlineNodeIds": ["node_1"],
            "status": "partial",
            "notes": "第四章已响应，实施章节待补。",
        }
    ]
    put = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "响应章节"}],
            "chapters": [{"id": "chap_1", "title": "响应章节"}],
            "responseMatrix": matrix,
        },
    )
    assert put.status_code == 200
    assert put.json()["responseMatrix"] == matrix

    got = client.get(f"/api/projects/{pid}/editor-state").json()
    assert got["responseMatrix"] == matrix


def test_response_matrix_partial_updates_do_not_clear_other_fields(client):
    pid = _create_project(client)
    matrix = [
        {
            "id": "mx_score_1",
            "kind": "scoring",
            "sourceKey": "scoring:总体架构",
            "sourceIndex": 0,
            "sourceText": "总体架构",
            "weight": "20%",
            "chapterIds": ["chap_arch"],
            "outlineNodeIds": ["node_arch"],
            "status": "covered",
            "notes": "已覆盖。",
        }
    ]

    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_arch", "title": "总体架构", "children": []}],
            "chapters": [{"id": "chap_arch", "title": "总体架构"}],
            "responseMatrix": matrix,
        },
    )

    facts = [{"id": "f1", "category": "要求", "content": "等保三级", "source": "manual"}]
    put_facts = client.put(
        f"/api/projects/{pid}/editor-state",
        json={"facts": facts},
    )
    assert put_facts.status_code == 200
    body = put_facts.json()
    assert body["responseMatrix"] == matrix
    assert body["outline"][0]["id"] == "node_arch"
    assert body["facts"] == facts

    put_null = client.put(
        f"/api/projects/{pid}/editor-state",
        json={"responseMatrix": None},
    )
    assert put_null.status_code == 200
    assert put_null.json()["responseMatrix"] == matrix

    put_matrix = client.put(
        f"/api/projects/{pid}/editor-state",
        json={"responseMatrix": []},
    )
    assert put_matrix.status_code == 200
    assert put_matrix.json()["responseMatrix"] == []
    assert put_matrix.json()["outline"][0]["id"] == "node_arch"


def test_response_matrix_normalizes_invalid_rows(client):
    pid = _create_project(client)
    put = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "chapters": [{"id": "chap_1", "title": "总体架构"}],
            "responseMatrix": [
                {"kind": "unknown", "sourceText": "非法行"},
                {
                    "kind": "scoring",
                    "sourceIndex": "2",
                    "sourceText": " 总体架构 ",
                    "status": "bad",
                    "chapterIds": ["chap_1", "chap_1", ""],
                    "outlineNodeIds": "not-list",
                },
            ]
        },
    )
    assert put.status_code == 200
    rows = put.json()["responseMatrix"]
    assert len(rows) == 1
    assert rows[0]["id"].startswith("mx_")
    assert rows[0]["kind"] == "scoring"
    assert rows[0]["sourceKey"] == "scoring:总体架构"
    assert rows[0]["sourceIndex"] == 2
    assert rows[0]["sourceText"] == "总体架构"
    assert rows[0]["chapterIds"] == ["chap_1"]
    assert rows[0]["outlineNodeIds"] == []
    assert rows[0]["status"] == "uncovered"


def test_response_matrix_reconciles_dead_outline_and_chapter_ids(client):
    pid = _create_project(client)
    matrix = [
        {
            "id": "mx_req_1",
            "kind": "requirement",
            "sourceKey": "requirement:等保三级",
            "sourceIndex": 0,
            "sourceText": "等保三级",
            "chapterIds": ["chap_a"],
            "outlineNodeIds": ["node_a"],
            "status": "covered",
            "notes": "",
        }
    ]
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_a", "title": "安全方案"}],
            "chapters": [{"id": "chap_a", "title": "安全方案"}],
            "responseMatrix": matrix,
        },
    ).json()
    assert saved["responseMatrix"][0]["status"] == "covered"

    changed = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_b", "title": "新安全方案"}],
            "chapters": [{"id": "chap_b", "title": "新安全方案"}],
        },
    ).json()
    row = changed["responseMatrix"][0]
    assert row["chapterIds"] == []
    assert row["outlineNodeIds"] == []
    assert row["status"] == "uncovered"


def test_response_matrix_bad_json_does_not_break_get(client):
    pid = _create_project(client)
    db = SessionLocal()
    try:
        db.add(
            ProjectEditorStateRow(
                project_id=pid,
                mode="ALIGNED",
                response_matrix_json="{bad json",
            )
        )
        db.commit()
    finally:
        db.close()

    got = client.get(f"/api/projects/{pid}/editor-state")
    assert got.status_code == 200
    assert got.json()["responseMatrix"] == []


def test_response_matrix_column_is_added_to_old_sqlite_table(tmp_path):
    old_engine = create_engine(f"sqlite:///{tmp_path / 'old.db'}")
    with old_engine.begin() as conn:
        conn.exec_driver_sql(
            """
            CREATE TABLE project_editor_states (
              project_id VARCHAR(64) PRIMARY KEY,
              mode VARCHAR(32) NOT NULL DEFAULT 'ALIGNED',
              updated_at DATETIME NOT NULL
            )
            """
        )

    ensure_schema_columns(old_engine)

    with old_engine.connect() as conn:
        columns = {
            row[1]
            for row in conn.execute(text("PRAGMA table_info(project_editor_states)"))
        }
    assert "response_matrix_json" in columns


def test_technical_export_includes_reconciled_response_matrix(client):
    """用途：验收技术标 Word 导出显示矩阵且不把失效关联当作覆盖。"""
    pid = _create_project(client)
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_security", "title": "安全设计"}],
            "chapters": [{"id": "chapter_security", "title": "安全实施"}],
            "responseMatrix": [
                {
                    "kind": "requirement",
                    "sourceText": "满足等保三级要求",
                    "chapterIds": ["chapter_security"],
                    "outlineNodeIds": ["node_security"],
                    "status": "covered",
                    "notes": "已在方案中说明。",
                }
            ],
        },
    )
    assert saved.status_code == 200

    # 模拟旧客户端或历史数据残留：导出路径必须独立收敛，而不是只信任 API 写入时的结果。
    db = SessionLocal()
    try:
        row = db.get(ProjectEditorStateRow, pid)
        assert row is not None
        row.response_matrix_json = json.dumps(
            [
                *saved.json()["responseMatrix"],
                {
                    "kind": "scoring",
                    "sourceText": "不存在的关联不得计为覆盖",
                    "weight": "10 分",
                    "chapterIds": ["chapter_missing"],
                    "outlineNodeIds": ["node_missing"],
                    "status": "covered",
                    "notes": "历史残留。",
                },
                {
                    "kind": "requirement",
                    "sourceText": "不响应项仍需留痕",
                    "status": "waived",
                    "notes": "不适用。",
                },
            ],
            ensure_ascii=False,
        )
        db.commit()
    finally:
        db.close()

    exported = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "export"},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(f"/api/projects/{pid}/export/download/{stored_name}")
    assert downloaded.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    assert any(paragraph.text == "六、响应矩阵" for paragraph in doc.paragraphs)
    matrix = doc.tables[-1]
    assert [cell.text for cell in matrix.rows[0].cells] == [
        "类型",
        "要求/评分点",
        "权重",
        "响应状态",
        "关联位置",
        "备注",
    ]
    table_text = "\n".join(cell.text for row in matrix.rows for cell in row.cells)
    assert "满足等保三级要求" in table_text
    assert "技术要求" in table_text
    assert "已覆盖" in table_text
    assert "正文：安全实施" in table_text
    assert "大纲：安全设计" in table_text
    assert "不存在的关联不得计为覆盖" in table_text
    assert "未覆盖" in table_text
    assert "chapter_missing" not in table_text
    assert "node_missing" not in table_text
    assert "不响应项仍需留痕" in table_text
    assert "不响应" in table_text


def test_business_export_does_not_include_response_matrix(client):
    """用途：验收响应矩阵仅属于技术标导出分支。"""
    project = client.post(
        "/api/projects", json={"name": "商务标响应矩阵隔离", "kind": "business"}
    ).json()
    pid = project["id"]
    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "responseMatrix": [
                {
                    "kind": "requirement",
                    "sourceText": "不应导出到商务标",
                    "status": "uncovered",
                }
            ]
        },
    )

    exported = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "export", "payload": {"mode": "business"}},
    )
    assert exported.status_code == 201
    stored_name = exported.json()["result"]["storedName"]
    downloaded = client.get(f"/api/projects/{pid}/export/download/{stored_name}")
    assert downloaded.status_code == 200

    from docx import Document  # type: ignore

    doc = Document(BytesIO(downloaded.content))
    body = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert not any(paragraph.text == "六、响应矩阵" for paragraph in doc.paragraphs)
    assert "不应导出到商务标" not in body


def test_response_match_returns_sanitized_suggestions_without_writing_state(
    client, monkeypatch
):
    """用途：验收智能建议只写任务结果，且过滤重复来源、非法 ID 与 waived 行。"""
    pid = _create_project(client)
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [
                {"id": "node_arch", "title": "总体架构"},
                {"id": "node_security", "title": "安全设计"},
            ],
            "chapters": [
                {"id": "chapter_arch", "title": "架构正文"},
                {"id": "chapter_security", "title": "安全正文"},
            ],
            "responseMatrix": [
                {
                    "kind": "requirement",
                    "sourceKey": "requirement:等保三级",
                    "sourceText": "等保三级",
                    "status": "uncovered",
                },
                {
                    "kind": "scoring",
                    "sourceKey": "scoring:总体架构",
                    "sourceText": "总体架构",
                    "weight": "20%",
                    "status": "partial",
                },
                {
                    "kind": "requirement",
                    "sourceKey": "requirement:不适用",
                    "sourceText": "不适用要求",
                    "status": "waived",
                },
            ],
        },
    )
    assert saved.status_code == 200
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        return ChatResult(
            content=json.dumps(
                [
                    {
                        "sourceKey": "requirement:等保三级",
                        "chapterIds": ["chapter_arch"],
                        "outlineNodeIds": ["node_arch"],
                        "status": "covered",
                        "confidence": 45,
                        "reason": "低置信度重复建议",
                    },
                    {
                        "sourceKey": "requirement:等保三级",
                        "chapterIds": ["chapter_security", "chapter_missing"],
                        "outlineNodeIds": ["node_security"],
                        "status": "covered",
                        "confidence": 91,
                        "reason": "安全章节直接回应等保要求",
                    },
                    {
                        "sourceKey": "scoring:总体架构",
                        "chapterIds": ["chapter_missing"],
                        "outlineNodeIds": ["node_missing"],
                        "status": "covered",
                        "confidence": 88,
                        "reason": "无效关联",
                    },
                    {
                        "sourceKey": "requirement:不适用",
                        "chapterIds": ["chapter_arch"],
                        "status": "covered",
                        "confidence": 99,
                    },
                    {
                        "sourceKey": "unknown:来源",
                        "chapterIds": ["chapter_arch"],
                        "status": "covered",
                        "confidence": 99,
                    },
                ],
                ensure_ascii=False,
            ),
            model="demo",
        )

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)
    response = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "response_match"},
    )
    assert response.status_code == 201
    task = response.json()
    assert task["status"] == "success"
    suggestions = task["result"]["suggestions"]
    assert len(suggestions) == 2
    requirement = suggestions[0]
    assert requirement["sourceKey"] == "requirement:等保三级"
    assert requirement["chapterIds"] == ["chapter_security"]
    assert requirement["outlineNodeIds"] == ["node_security"]
    assert requirement["status"] == "covered"
    assert requirement["confidence"] == 91
    assert requirement["base"] == {
        "chapterIds": [],
        "outlineNodeIds": [],
        "status": "uncovered",
    }
    invalid_links = suggestions[1]
    assert invalid_links["sourceKey"] == "scoring:总体架构"
    assert invalid_links["chapterIds"] == []
    assert invalid_links["outlineNodeIds"] == []
    assert invalid_links["status"] == "uncovered"
    assert task["result"]["skippedInvalidCount"] == 5
    # 默认 batch0 兼容：旧客户端不传 payload 时语义不变，并返回批次元数据
    assert task["result"]["candidateBatchIndex"] == 0
    assert task["result"]["candidateBatchCount"] == 1
    assert task["result"]["isLastCandidateBatch"] is True
    assert task["result"]["sourceBatchIndex"] == 0
    assert task["result"]["sourceBatchCount"] == 1
    assert task["result"]["isLastSourceBatch"] is True
    assert task["result"]["chapterCandidateTotal"] == 2
    assert task["result"]["outlineCandidateTotal"] == 2
    assert task["result"]["chapterCandidateInBatch"] == 2
    assert task["result"]["outlineCandidateInBatch"] == 2
    assert task["result"]["sourceCount"] == 2
    assert task["result"]["totalSourceCount"] == 2
    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def _extract_prompt_option_ids(user_content: str, section: str) -> list[str]:
    """用途：从模型 user 提示中解析章节/大纲候选 id，用于验收分批窗口。"""
    marker = f"【{section}】\n"
    start = user_content.find(marker)
    if start < 0:
        return []
    rest = user_content[start + len(marker) :]
    end = rest.find("\n\n【")
    block = rest if end < 0 else rest[:end]
    if block.strip() == "无":
        return []
    ids: list[str] = []
    for line in block.splitlines():
        line = line.strip()
        if not line.startswith("- id="):
            continue
        part = line[len("- id=") :]
        option_id = part.split("；", 1)[0].strip()
        if option_id:
            ids.append(option_id)
    return ids


def _extract_prompt_source_keys(user_content: str) -> list[str]:
    """用途：从模型 user 提示中解析待匹配条目的 sourceKey，用于验收来源分页窗口。"""
    marker = "【待匹配条目】\n"
    start = user_content.find(marker)
    if start < 0:
        return []
    rest = user_content[start + len(marker) :]
    end = rest.find("\n\n【")
    block = rest if end < 0 else rest[:end]
    keys: list[str] = []
    for line in block.splitlines():
        line = line.strip()
        if not line.startswith("- sourceKey="):
            continue
        part = line[len("- sourceKey=") :]
        source_key = part.split("；", 1)[0].strip()
        if source_key:
            keys.append(source_key)
    return keys


def test_response_match_candidate_batches_cover_without_overlap(client, monkeypatch):
    """用途：>120 章 / >160 大纲时批窗口无重叠并覆盖全量；任意批不写 responseMatrix。"""
    pid = _create_project(client)
    chapters = [
        {"id": f"chapter_{i:03d}", "title": f"章节{i:03d}"} for i in range(250)
    ]
    outline = [{"id": f"node_{i:03d}", "title": f"大纲{i:03d}"} for i in range(200)]
    # 构造超过来源 80 上限，验证 sourceCount 仍截断
    matrix = [
        {
            "kind": "requirement",
            "sourceKey": f"requirement:来源{i:03d}",
            "sourceText": f"来源{i:03d}",
            "status": "uncovered",
        }
        for i in range(90)
    ]
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": outline,
            "chapters": chapters,
            "responseMatrix": matrix,
        },
    )
    assert saved.status_code == 200
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]

    captured: list[str] = []

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        user_content = messages[-1]["content"]
        captured.append(user_content)
        chapter_ids = _extract_prompt_option_ids(user_content, "章节候选")
        outline_ids = _extract_prompt_option_ids(user_content, "大纲候选")
        # 故意混入跨批非法 id，确认仅本批 ID 可通过规范化
        foreign_chapter = "chapter_200" if "chapter_000" in chapter_ids else "chapter_000"
        return ChatResult(
            content=json.dumps(
                [
                    {
                        "sourceKey": "requirement:来源000",
                        "chapterIds": ([chapter_ids[0]] if chapter_ids else [])
                        + [foreign_chapter],
                        "outlineNodeIds": ([outline_ids[0]] if outline_ids else []),
                        "status": "covered",
                        "confidence": 80 + len(captured),
                        "reason": f"批次探测{len(captured)}",
                    }
                ],
                ensure_ascii=False,
            ),
            model="demo",
        )

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)

    batch0 = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "response_match", "payload": {"candidateBatchIndex": 0}},
    ).json()
    batch1 = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "response_match", "payload": {"candidateBatchIndex": 1}},
    ).json()
    batch2 = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "response_match", "payload": {"candidateBatchIndex": 2}},
    ).json()

    assert batch0["status"] == "success"
    assert batch1["status"] == "success"
    assert batch2["status"] == "success"

    r0, r1, r2 = batch0["result"], batch1["result"], batch2["result"]
    assert r0["candidateBatchCount"] == 3
    assert r1["candidateBatchCount"] == 3
    assert r2["candidateBatchCount"] == 3
    assert r0["candidateBatchIndex"] == 0
    assert r1["candidateBatchIndex"] == 1
    assert r2["candidateBatchIndex"] == 2
    assert r0["isLastCandidateBatch"] is False
    assert r1["isLastCandidateBatch"] is False
    assert r2["isLastCandidateBatch"] is True
    assert r0["chapterCandidateTotal"] == 250
    assert r0["outlineCandidateTotal"] == 200
    assert r0["chapterCandidateInBatch"] == 120
    assert r0["outlineCandidateInBatch"] == 160
    assert r1["chapterCandidateInBatch"] == 120
    assert r1["outlineCandidateInBatch"] == 40
    assert r2["chapterCandidateInBatch"] == 10
    assert r2["outlineCandidateInBatch"] == 0
    # 未传 sourceBatchIndex 时等价来源页 0；单次仍最多 80 条来源
    assert r0["sourceCount"] == 80
    assert r0["totalSourceCount"] == 90
    assert r0["sourceBatchIndex"] == 0
    assert r0["sourceBatchCount"] == 2
    assert r0["isLastSourceBatch"] is False
    assert r1["sourceCount"] == 80
    assert r2["sourceCount"] == 80

    c0 = _extract_prompt_option_ids(captured[0], "章节候选")
    c1 = _extract_prompt_option_ids(captured[1], "章节候选")
    c2 = _extract_prompt_option_ids(captured[2], "章节候选")
    o0 = _extract_prompt_option_ids(captured[0], "大纲候选")
    o1 = _extract_prompt_option_ids(captured[1], "大纲候选")
    o2 = _extract_prompt_option_ids(captured[2], "大纲候选")
    assert len(c0) == 120 and len(c1) == 120 and len(c2) == 10
    assert set(c0).isdisjoint(c1) and set(c0).isdisjoint(c2) and set(c1).isdisjoint(c2)
    assert set(c0) | set(c1) | set(c2) == {f"chapter_{i:03d}" for i in range(250)}
    assert len(o0) == 160 and len(o1) == 40 and len(o2) == 0
    assert set(o0).isdisjoint(o1)
    assert set(o0) | set(o1) == {f"node_{i:03d}" for i in range(200)}

    # 本批 ID 保留，跨批 foreign 被剔除
    assert r0["suggestions"][0]["chapterIds"] == ["chapter_000"]
    assert r1["suggestions"][0]["chapterIds"] == ["chapter_120"]
    assert r2["suggestions"][0]["chapterIds"] == ["chapter_240"]

    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def test_response_match_source_batches_page_without_overflow(client, monkeypatch):
    """用途：81 条非 waived 来源分页；页 0=80、页 1=1；prompt 不含跨页来源。"""
    pid = _create_project(client)
    matrix = [
        {
            "kind": "requirement",
            "sourceKey": f"requirement:来源{i:03d}",
            "sourceText": f"来源{i:03d}",
            "status": "uncovered",
        }
        for i in range(81)
    ]
    # waived 不计入分页总数
    matrix.append(
        {
            "kind": "requirement",
            "sourceKey": "requirement:已放弃",
            "sourceText": "已放弃",
            "status": "waived",
        }
    )
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "大纲1"}],
            "chapters": [{"id": "chapter_1", "title": "章节1"}],
            "responseMatrix": matrix,
        },
    )
    assert saved.status_code == 200
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    captured: list[str] = []

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        user_content = messages[-1]["content"]
        captured.append(user_content)
        keys = _extract_prompt_source_keys(user_content)
        return ChatResult(
            content=json.dumps(
                [
                    {
                        "sourceKey": keys[0],
                        "chapterIds": ["chapter_1"],
                        "outlineNodeIds": ["node_1"],
                        "status": "covered",
                        "confidence": 80 + len(captured),
                        "reason": f"来源页探测{len(captured)}",
                    }
                ],
                ensure_ascii=False,
            ),
            model="demo",
        )

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)

    page0 = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={
            "type": "response_match",
            "payload": {"sourceBatchIndex": 0, "candidateBatchIndex": 0},
        },
    ).json()
    page1 = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={
            "type": "response_match",
            "payload": {"sourceBatchIndex": 1, "candidateBatchIndex": 0},
        },
    ).json()

    assert page0["status"] == "success"
    assert page1["status"] == "success"
    r0, r1 = page0["result"], page1["result"]
    assert r0["sourceBatchIndex"] == 0
    assert r1["sourceBatchIndex"] == 1
    assert r0["sourceBatchCount"] == 2
    assert r1["sourceBatchCount"] == 2
    assert r0["sourceCount"] == 80
    assert r1["sourceCount"] == 1
    assert r0["totalSourceCount"] == 81
    assert r1["totalSourceCount"] == 81
    assert r0["isLastSourceBatch"] is False
    assert r1["isLastSourceBatch"] is True
    assert r0["candidateBatchIndex"] == 0
    assert r0["isLastCandidateBatch"] is True

    keys0 = _extract_prompt_source_keys(captured[0])
    keys1 = _extract_prompt_source_keys(captured[1])
    assert len(keys0) == 80
    assert len(keys1) == 1
    assert keys0 == [f"requirement:来源{i:03d}" for i in range(80)]
    assert keys1 == ["requirement:来源080"]
    assert "requirement:来源080" not in keys0
    assert "requirement:已放弃" not in keys0
    assert "requirement:已放弃" not in keys1
    assert set(keys0).isdisjoint(set(keys1))

    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def test_response_match_source_and_candidate_nested_batches(client, monkeypatch):
    """用途：来源页 × 候选批嵌套；每请求仅含对应来源页与候选窗口，ID 校验不放宽。"""
    pid = _create_project(client)
    chapters = [
        {"id": f"chapter_{i:03d}", "title": f"章节{i:03d}"} for i in range(130)
    ]
    outline = [{"id": f"node_{i:03d}", "title": f"大纲{i:03d}"} for i in range(10)]
    matrix = [
        {
            "kind": "requirement",
            "sourceKey": f"requirement:来源{i:03d}",
            "sourceText": f"来源{i:03d}",
            "status": "uncovered",
        }
        for i in range(85)
    ]
    saved = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": outline,
            "chapters": chapters,
            "responseMatrix": matrix,
        },
    )
    assert saved.status_code == 200
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    captured: list[tuple[int, int, str]] = []

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        user_content = messages[-1]["content"]
        source_keys = _extract_prompt_source_keys(user_content)
        chapter_ids = _extract_prompt_option_ids(user_content, "章节候选")
        # 从 sourceKey 推断来源页：0..79 为页0，80+ 为页1
        src_idx = 0 if source_keys and source_keys[0].endswith("000") else 1
        cand_idx = 0 if chapter_ids and chapter_ids[0] == "chapter_000" else 1
        captured.append((src_idx, cand_idx, user_content))
        foreign_chapter = "chapter_120" if "chapter_000" in chapter_ids else "chapter_000"
        return ChatResult(
            content=json.dumps(
                [
                    {
                        "sourceKey": source_keys[0],
                        "chapterIds": ([chapter_ids[0]] if chapter_ids else [])
                        + [foreign_chapter],
                        "outlineNodeIds": [],
                        "status": "covered",
                        "confidence": 75,
                        "reason": f"嵌套探测 s{src_idx}c{cand_idx}",
                    }
                ],
                ensure_ascii=False,
            ),
            model="demo",
        )

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)

    results = []
    for source_idx in (0, 1):
        for cand_idx in (0, 1):
            task = client.post(
                f"/api/projects/{pid}/tasks?sync=true",
                json={
                    "type": "response_match",
                    "payload": {
                        "sourceBatchIndex": source_idx,
                        "candidateBatchIndex": cand_idx,
                    },
                },
            ).json()
            assert task["status"] == "success"
            results.append(task["result"])

    assert len(captured) == 4
    assert len(results) == 4
    for result in results:
        assert result["sourceBatchCount"] == 2
        assert result["candidateBatchCount"] == 2
        assert result["totalSourceCount"] == 85

    r_s0c0, r_s0c1, r_s1c0, r_s1c1 = results
    assert r_s0c0["sourceBatchIndex"] == 0 and r_s0c0["candidateBatchIndex"] == 0
    assert r_s0c0["sourceCount"] == 80
    assert r_s0c0["isLastSourceBatch"] is False
    assert r_s0c0["isLastCandidateBatch"] is False
    assert r_s0c1["sourceCount"] == 80
    assert r_s0c1["isLastCandidateBatch"] is True
    assert r_s1c0["sourceBatchIndex"] == 1
    assert r_s1c0["sourceCount"] == 5
    assert r_s1c0["isLastSourceBatch"] is True
    assert r_s1c1["isLastSourceBatch"] is True
    assert r_s1c1["isLastCandidateBatch"] is True

    keys_s0 = _extract_prompt_source_keys(captured[0][2])
    keys_s1 = _extract_prompt_source_keys(captured[2][2])
    assert len(keys_s0) == 80
    assert len(keys_s1) == 5
    assert keys_s0 == [f"requirement:来源{i:03d}" for i in range(80)]
    assert keys_s1 == [f"requirement:来源{i:03d}" for i in range(80, 85)]
    assert set(keys_s0).isdisjoint(set(keys_s1))

    # 同一来源页上候选窗口无重叠；跨来源页可重复同一候选窗口
    c_s0c0 = _extract_prompt_option_ids(captured[0][2], "章节候选")
    c_s0c1 = _extract_prompt_option_ids(captured[1][2], "章节候选")
    c_s1c0 = _extract_prompt_option_ids(captured[2][2], "章节候选")
    assert len(c_s0c0) == 120 and len(c_s0c1) == 10
    assert set(c_s0c0).isdisjoint(c_s0c1)
    assert c_s1c0 == c_s0c0

    # 跨批 foreign chapter 仍被剔除
    assert r_s0c0["suggestions"][0]["chapterIds"] == ["chapter_000"]
    assert r_s0c1["suggestions"][0]["chapterIds"] == ["chapter_120"]
    assert r_s1c0["suggestions"][0]["chapterIds"] == ["chapter_000"]
    assert r_s1c1["suggestions"][0]["chapterIds"] == ["chapter_120"]

    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def test_response_match_batch_out_of_range_fails(client, monkeypatch):
    """用途：候选批次越界时任务失败，且不改 editor-state。"""
    pid = _create_project(client)
    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "大纲1"}],
            "chapters": [{"id": "chapter_1", "title": "章节1"}],
            "responseMatrix": [
                {
                    "kind": "requirement",
                    "sourceKey": "requirement:一条",
                    "sourceText": "一条",
                    "status": "uncovered",
                }
            ],
        },
    )
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    called = {"n": 0}

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        called["n"] += 1
        return ChatResult(content="[]", model="demo")

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)
    response = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={"type": "response_match", "payload": {"candidateBatchIndex": 9}},
    )
    assert response.status_code == 201
    task = response.json()
    assert task["status"] == "failed"
    detail = (task.get("error") or task.get("message") or "")
    assert "候选批次越界" in detail
    assert called["n"] == 0
    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def test_response_match_source_batch_out_of_range_fails(client, monkeypatch):
    """用途：来源页越界时任务 failed、模型 0 次、editor-state 不变。"""
    pid = _create_project(client)
    matrix = [
        {
            "kind": "requirement",
            "sourceKey": f"requirement:来源{i:03d}",
            "sourceText": f"来源{i:03d}",
            "status": "uncovered",
        }
        for i in range(81)
    ]
    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "大纲1"}],
            "chapters": [{"id": "chapter_1", "title": "章节1"}],
            "responseMatrix": matrix,
        },
    )
    before = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    called = {"n": 0}

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        called["n"] += 1
        return ChatResult(content="[]", model="demo")

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)
    response = client.post(
        f"/api/projects/{pid}/tasks?sync=true",
        json={
            "type": "response_match",
            "payload": {"sourceBatchIndex": 9, "candidateBatchIndex": 0},
        },
    )
    assert response.status_code == 201
    task = response.json()
    assert task["status"] == "failed"
    detail = (task.get("error") or task.get("message") or "")
    assert "来源批次越界" in detail
    assert called["n"] == 0
    after = client.get(f"/api/projects/{pid}/editor-state").json()["responseMatrix"]
    assert after == before


def test_response_match_invalid_batch_index_defaults_to_zero(client, monkeypatch):
    """用途：缺失/非法类型/负值 candidateBatchIndex 与 sourceBatchIndex 均视为 0。"""
    pid = _create_project(client)
    client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "大纲1"}],
            "chapters": [{"id": "chapter_1", "title": "章节1"}],
            "responseMatrix": [
                {
                    "kind": "requirement",
                    "sourceKey": "requirement:一条",
                    "sourceText": "一条",
                    "status": "uncovered",
                }
            ],
        },
    )

    def fake_chat(db, workspace_id, *, messages, temperature=0.4, timeout_sec=120.0):
        return ChatResult(
            content=json.dumps(
                [
                    {
                        "sourceKey": "requirement:一条",
                        "chapterIds": ["chapter_1"],
                        "outlineNodeIds": ["node_1"],
                        "status": "covered",
                        "confidence": 70,
                        "reason": "默认批",
                    }
                ],
                ensure_ascii=False,
            ),
            model="demo",
        )

    monkeypatch.setattr("app.services.llm_service.chat_completion", fake_chat)
    for payload in (
        None,
        {},
        {"candidateBatchIndex": "bad"},
        {"candidateBatchIndex": -3},
        {"candidateBatchIndex": True},
        {"candidateBatchIndex": 1.5},
        {"candidateBatchIndex": "1"},
        {"sourceBatchIndex": "bad"},
        {"sourceBatchIndex": -2},
        {"sourceBatchIndex": True},
        {"sourceBatchIndex": 1.5},
        {"sourceBatchIndex": "1"},
        {"sourceBatchIndex": None, "candidateBatchIndex": None},
    ):
        body: dict = {"type": "response_match"}
        if payload is not None:
            body["payload"] = payload
        task = client.post(
            f"/api/projects/{pid}/tasks?sync=true",
            json=body,
        ).json()
        assert task["status"] == "success"
        assert task["result"]["candidateBatchIndex"] == 0
        assert task["result"]["isLastCandidateBatch"] is True
        assert task["result"]["sourceBatchIndex"] == 0
        assert task["result"]["isLastSourceBatch"] is True


def test_response_matrix_version_stable_and_empty(client):
    """用途：空矩阵也有稳定版本；改概述不改矩阵版本。"""
    pid = _create_project(client)
    first = client.get(f"/api/projects/{pid}/editor-state").json()
    assert first["responseMatrix"] == []
    assert isinstance(first["responseMatrixVersion"], str)
    assert first["responseMatrixVersion"].startswith("rmv_")
    empty_version = first["responseMatrixVersion"]

    second = client.get(f"/api/projects/{pid}/editor-state").json()
    assert second["responseMatrixVersion"] == empty_version

    matrix = [
        {
            "id": "mx_req_v1",
            "kind": "requirement",
            "sourceKey": "requirement:等保三级",
            "sourceIndex": 0,
            "sourceText": "等保三级",
            "chapterIds": ["chap_1"],
            "outlineNodeIds": ["node_1"],
            "status": "partial",
            "notes": "初稿",
        }
    ]
    put = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "outline": [{"id": "node_1", "title": "安全"}],
            "chapters": [{"id": "chap_1", "title": "安全"}],
            "responseMatrix": matrix,
            "responseMatrixVersion": empty_version,
        },
    )
    assert put.status_code == 200
    matrix_version = put.json()["responseMatrixVersion"]
    assert matrix_version != empty_version
    assert put.json()["responseMatrix"][0]["sourceText"] == "等保三级"

    overview_only = client.put(
        f"/api/projects/{pid}/editor-state",
        json={"analysisOverview": "仅改概述不应改变矩阵版本"},
    )
    assert overview_only.status_code == 200
    assert overview_only.json()["responseMatrixVersion"] == matrix_version
    assert overview_only.json()["analysisOverview"] == "仅改概述不应改变矩阵版本"


def test_response_matrix_version_match_save_and_stale_conflict(client):
    """用途：版本匹配可保存；两客户端陈旧版本 409 且远端矩阵不变。"""
    pid = _create_project(client)
    base = client.get(f"/api/projects/{pid}/editor-state").json()
    v0 = base["responseMatrixVersion"]

    client_a_matrix = [
        {
            "id": "mx_a",
            "kind": "requirement",
            "sourceKey": "requirement:A端写入",
            "sourceIndex": 0,
            "sourceText": "A端写入",
            "chapterIds": [],
            "outlineNodeIds": [],
            "status": "uncovered",
            "notes": "A",
        }
    ]
    a_ok = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "responseMatrix": client_a_matrix,
            "responseMatrixVersion": v0,
        },
    )
    assert a_ok.status_code == 200
    assert a_ok.json()["responseMatrix"][0]["sourceText"] == "A端写入"
    v_a = a_ok.json()["responseMatrixVersion"]

    client_b_stale = [
        {
            "id": "mx_b",
            "kind": "requirement",
            "sourceKey": "requirement:B端陈旧",
            "sourceIndex": 0,
            "sourceText": "B端陈旧",
            "chapterIds": [],
            "outlineNodeIds": [],
            "status": "uncovered",
            "notes": "B",
        }
    ]
    conflict = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "responseMatrix": client_b_stale,
            "responseMatrixVersion": v0,
            "analysisOverview": "冲突时不得写入概述",
        },
    )
    assert conflict.status_code == 409
    detail = conflict.json()["detail"]
    assert "message" in detail and detail["message"]
    assert detail["currentResponseMatrixVersion"] == v_a
    assert detail["responseMatrix"][0]["sourceText"] == "A端写入"

    after = client.get(f"/api/projects/{pid}/editor-state").json()
    assert after["responseMatrix"][0]["sourceText"] == "A端写入"
    assert after["responseMatrixVersion"] == v_a
    assert after.get("analysisOverview") in (None, "", after.get("analysis", {}).get("overview"))
    # 冲突整包不写：概述不应被陈旧请求改掉
    assert (after.get("analysisOverview") or "") != "冲突时不得写入概述"

    matched = client.put(
        f"/api/projects/{pid}/editor-state",
        json={
            "responseMatrix": client_b_stale,
            "responseMatrixVersion": v_a,
        },
    )
    assert matched.status_code == 200
    assert matched.json()["responseMatrix"][0]["sourceText"] == "B端陈旧"


def test_response_matrix_legacy_put_without_version(client):
    """用途：旧客户端不带 responseMatrixVersion 仍可写矩阵。"""
    pid = _create_project(client)
    matrix = [
        {
            "kind": "scoring",
            "sourceText": "旧客户端写入",
            "weight": "10%",
            "chapterIds": [],
            "outlineNodeIds": [],
            "status": "uncovered",
        }
    ]
    put = client.put(
        f"/api/projects/{pid}/editor-state",
        json={"responseMatrix": matrix},
    )
    assert put.status_code == 200
    assert put.json()["responseMatrix"][0]["sourceText"] == "旧客户端写入"
    assert put.json()["responseMatrixVersion"].startswith("rmv_")


def test_response_matrix_concurrent_versioned_puts_one_wins(client):
    """
    用途：独立 Session 并发带同一 expected version 的矩阵 PUT，恰好一方成功一方 409。
    对接：editor_state_service 写锁；不得用顺序两次调用代替。
    """
    import threading
    from concurrent.futures import ThreadPoolExecutor

    from app.core.database import SessionLocal
    from app.services import editor_state_service
    from app.services.editor_state_service import ResponseMatrixVersionConflict

    pid = _create_project(client)
    base = client.get(f"/api/projects/{pid}/editor-state").json()
    v0 = base["responseMatrixVersion"]
    workspace_id = "ws_local"

    def _matrix(label: str) -> list[dict]:
        return [
            {
                "id": f"mx_{label}",
                "kind": "requirement",
                "sourceKey": f"requirement:{label}",
                "sourceIndex": 0,
                "sourceText": label,
                "chapterIds": [],
                "outlineNodeIds": [],
                "status": "uncovered",
                "notes": label,
            }
        ]

    barrier = threading.Barrier(2)
    outcomes: list[tuple[str, str | None]] = []

    def worker(label: str) -> tuple[str, str | None]:
        db = SessionLocal()
        try:
            barrier.wait(timeout=5)
            try:
                data = editor_state_service.upsert_editor_state(
                    db,
                    workspace_id,
                    pid,
                    response_matrix=_matrix(label),
                    response_matrix_version=v0,
                )
                return ("ok", data["responseMatrix"][0]["sourceText"])
            except ResponseMatrixVersionConflict as exc:
                # 冲突方必须看到获胜者矩阵
                winner = (
                    exc.current_matrix[0]["sourceText"]
                    if exc.current_matrix
                    else None
                )
                return ("conflict", winner)
        finally:
            db.close()

    with ThreadPoolExecutor(max_workers=2) as pool:
        futures = [pool.submit(worker, "并发甲"), pool.submit(worker, "并发乙")]
        outcomes = [f.result(timeout=15) for f in futures]

    statuses = sorted(o[0] for o in outcomes)
    assert statuses == ["conflict", "ok"], outcomes
    ok_label = next(o[1] for o in outcomes if o[0] == "ok")
    conflict_seen = next(o[1] for o in outcomes if o[0] == "conflict")
    assert ok_label in ("并发甲", "并发乙")
    assert conflict_seen == ok_label

    final = client.get(f"/api/projects/{pid}/editor-state").json()
    assert final["responseMatrix"][0]["sourceText"] == ok_label
    assert final["responseMatrixVersion"] != v0
