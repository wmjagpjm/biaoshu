"""
模块：Word 导出服务
用途：将 editor-state 导出为 .docx，应用 workspace 默认 ExportFormat：
  - 纸张/边距/页眉页脚/页码
  - 正文与标题字体色间距
  - 标题段落边框与分级底色（heading_border）
  - 标题编号（numbering_format / numbering_template）
  - 列表符号（list_style / ordered_list_style）
  - 评分点与 Markdown 表格（table.* 边框/表头/首列）
  - 章节 body 粗解析 Markdown（列表/表格/小标题）
  - 技术标响应矩阵（仅输出当前有效的章节/大纲关联）
对接：export 任务；settings.export_format_json；前端 ExportFormatConfig
二次开发：
  - 正文图片仅支持项目内 biaoshu-image://file_<16位十六进制> 独占行引用
  - 字段兼容 snake_case 与 camelCase（_g）
  - min_heading_left_enabled：仅叶子标题左侧强调线；structure 仍未接线，不得误称整章页框
"""

from __future__ import annotations

import io
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.entities import ProjectEditorStateRow
from app.services.project_service import get_project
from app.services import file_service, settings_service


def _loads(raw: str | None) -> Any:
    if not raw:
        return None
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return None


# 中文字号 → 磅值
_CN_SIZE_PT: dict[str, float] = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
}

# 纸张 mm（宽×高，纵向）
_PAPER_MM: dict[str, tuple[float, float]] = {
    "a4": (210, 297),
    "a3": (297, 420),
    "a5": (148, 210),
    "b4": (250, 353),
    "b5": (176, 250),
    "letter": (215.9, 279.4),
    "legal": (215.9, 355.6),
    "16k": (184, 260),
}


def _g(d: dict | None, *keys: str, default: Any = None) -> Any:
    """用途：多键名兼容取值（snake / camel）。"""
    if not isinstance(d, dict):
        return default
    for k in keys:
        if k in d and d[k] is not None:
            return d[k]
    return default


def _pt(size_name: str | None, default: float = 12.0) -> float:
    if not size_name:
        return default
    if size_name in _CN_SIZE_PT:
        return _CN_SIZE_PT[size_name]
    try:
        return float(str(size_name).replace("pt", "").strip())
    except ValueError:
        return default


def _align(name: str | None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    m = {
        "居中对齐": WD_ALIGN_PARAGRAPH.CENTER,
        "居中": WD_ALIGN_PARAGRAPH.CENTER,
        "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
        "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return m.get(name or "", WD_ALIGN_PARAGRAPH.LEFT)


def _rgb(color: str | None):
    """用途：#RRGGBB → RGBColor；失败返回 None。"""
    if not color or not isinstance(color, str):
        return None
    s = color.strip().lstrip("#")
    if len(s) == 3:
        s = "".join(c * 2 for c in s)
    if len(s) != 6 or not re.fullmatch(r"[0-9a-fA-F]{6}", s):
        return None
    from docx.shared import RGBColor  # type: ignore

    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


# ── 标题编号（对齐前端 ExportFormat headings[].numbering_*）────────────

_ZH_DIGITS = "零一二三四五六七八九"


def _to_zh(n: int) -> str:
    """用途：阿拉伯数字 → 中文数字（1～99 常用）。"""
    if n <= 0:
        return str(n)
    if n < 10:
        return _ZH_DIGITS[n]
    if n == 10:
        return "十"
    if n < 20:
        return "十" + _ZH_DIGITS[n - 10]
    if n < 100:
        tens, ones = divmod(n, 10)
        return _ZH_DIGITS[tens] + "十" + (_ZH_DIGITS[ones] if ones else "")
    return str(n)


def _to_circled(n: int) -> str:
    """用途：1–20 → ①–⑳；超出回退阿拉伯数字。"""
    if 1 <= n <= 20:
        return chr(0x2460 + n - 1)
    return str(n)


# 标题已自带编号时不再前缀（避免「第一章 第一章 xxx」）
_ALREADY_NUMBERED = re.compile(
    r"^(?:"
    r"第[一二三四五六七八九十百千零〇\d]+[章节条款部篇]"
    r"|[（(]?\d+(?:\.\d+)*[)）．.、]"
    r"|[一二三四五六七八九十]+[、．.]"
    r")"
)


class HeadingNumberer:
    """
    用途：多级标题计数 + 按 numbering_format/template 生成前缀。
    对接：export_service 大纲 walk / 正文章标题。
    """

    def __init__(self, headings_cfg: list | None = None) -> None:
        self.headings = headings_cfg if isinstance(headings_cfg, list) else []
        self.counts = [0] * 6

    def next_prefix(self, level: int) -> str:
        """
        level: 0 起（0=一级章，1=二级…）。
        返回编号前缀字符串（不含标题正文）；无法编号时返回 \"\"。
        """
        level = max(0, min(int(level), 5))
        self.counts[level] += 1
        for i in range(level + 1, 6):
            self.counts[i] = 0
        path = [c for c in self.counts[: level + 1]]
        hcfg = self.headings[level] if level < len(self.headings) else None
        if not isinstance(hcfg, dict):
            hcfg = {}
        return format_heading_number(level, path, hcfg)


def format_heading_number(
    level: int, path: list[int], hcfg: dict | None
) -> str:
    """
    用途：根据编号格式/模板生成前缀。
    path: 各级当前序号，如 [1,2,1] 表示 1.2.1。
    """
    hcfg = hcfg if isinstance(hcfg, dict) else {}
    fmt = str(
        _g(hcfg, "numbering_format", "numberingFormat", default="outline-decimal")
        or "outline-decimal"
    ).strip()
    tpl = str(
        _g(hcfg, "numbering_template", "numberingTemplate", default="") or ""
    ).strip()

    if not path:
        return ""

    num = path[level] if level < len(path) else path[-1]
    full = ".".join(str(x) for x in path)

    if fmt == "custom" and tpl:
        s = tpl
        s = s.replace("{zh}", _to_zh(num))
        s = s.replace("{num}", str(num))
        s = s.replace("{full}", full)
        s = s.replace("{tail}", full)
        s = s.replace("{circled}", _to_circled(num))
        for i in range(1, 7):
            part = ".".join(str(x) for x in path[:i]) if len(path) >= i else full
            s = s.replace(f"{{tail{i}}}", part)
        return s.strip()

    # outline-decimal 或未配置 custom 模板
    return full


def compose_heading_text(title: str, prefix: str) -> str:
    """用途：前缀 + 标题；已有编号则原样返回。"""
    title = (title or "").strip() or "未命名"
    prefix = (prefix or "").strip()
    if not prefix:
        return title
    if _ALREADY_NUMBERED.match(title):
        return title
    if title.startswith(prefix):
        return title
    # 「第×章」类与标题之间加空格，便于阅读
    return f"{prefix} {title}"


# ── 列表 / 表格（对齐 body_text / table 配置）────────────────────────

_LIST_BULLETS: dict[str, str] = {
    "none": "",
    "disc": "•",
    "circle": "○",
    "square": "■",
    "diamond": "◆",
    "dash": "–",
    "check": "✓",
    "arrow": "➢",
    "sparkle": "✧",
}


def _bullet_prefix(list_style: str | None) -> str:
    return _LIST_BULLETS.get(str(list_style or "disc"), "•")


def _ordered_prefix(style: str | None, index: int) -> str:
    """index 从 1 起。"""
    s = str(style or "decimal-dot")
    if s == "decimal-dot":
        return f"{index}."
    if s == "decimal-paren":
        return f"{index})"
    if s == "decimal-full-paren":
        return f"（{index}）"
    if s == "chinese-dot":
        return f"{_to_zh(index)}、"
    if s == "chinese-paren":
        return f"（{_to_zh(index)}）"
    if s == "lower-alpha":
        return f"{chr(ord('a') + (index - 1) % 26)}."
    if s == "upper-alpha":
        return f"{chr(ord('A') + (index - 1) % 26)}."
    if s == "lower-roman":
        romans = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x"]
        return f"{romans[(index - 1) % len(romans)]}."
    if s == "upper-roman":
        romans = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
        return f"{romans[(index - 1) % len(romans)]}."
    return f"{index}."


def _body_list_cfg(template: dict | None) -> dict:
    body = {}
    if isinstance(template, dict):
        body = _g(template, "body_text", "bodyText", default={}) or {}
    if not isinstance(body, dict):
        body = {}
    return {
        "list_style": str(
            _g(body, "list_style", "listStyle", default="disc") or "disc"
        ),
        "ordered_list_style": str(
            _g(body, "ordered_list_style", "orderedListStyle", default="decimal-dot")
            or "decimal-dot"
        ),
        "list_indent_chars": float(
            _g(body, "list_indent_chars", "listIndentChars", default=2) or 2
        ),
    }


def _table_cfg(template: dict | None) -> dict:
    if not isinstance(template, dict):
        return {}
    t = _g(template, "table", default={}) or {}
    return t if isinstance(t, dict) else {}


def _image_cfg(template: dict | None) -> dict:
    """用途：读取正文图片样式，并将不可信宽度限制在可用版心比例内。"""
    raw = _g(template, "image", default={}) if isinstance(template, dict) else {}
    if not isinstance(raw, dict):
        raw = {}
    try:
        max_width = float(
            _g(raw, "max_width_percent", "maxWidthPercent", default=90) or 90
        )
    except (TypeError, ValueError):
        max_width = 90
    return {
        "max_width_percent": max(20, min(100, max_width)),
        "alignment": str(_g(raw, "alignment", default="居中对齐") or "居中对齐"),
        "caption_font": str(
            _g(raw, "caption_font", "captionFont", default="宋体") or "宋体"
        ),
        "caption_size": str(
            _g(raw, "caption_size", "captionSize", default="小五") or "小五"
        ),
        "caption_alignment": str(
            _g(raw, "caption_alignment", "captionAlignment", default="居中对齐")
            or "居中对齐"
        ),
        "caption_bold": _strict_bool(
            _g(raw, "caption_bold", "captionBold", default=False)
        ),
        "caption_italic": _strict_bool(
            _g(raw, "caption_italic", "captionItalic", default=False)
        ),
    }


def _add_list_item(
    doc,
    text: str,
    *,
    ordered: bool = False,
    index: int = 1,
    list_cfg: dict | None = None,
) -> None:
    """用途：按模板前缀写列表项（不用 Word 内置 List Bullet，以便自定义符号）。"""
    from docx.shared import Twips  # type: ignore

    cfg = list_cfg or {}
    text = str(text).strip()
    if not text:
        return
    if ordered:
        prefix = _ordered_prefix(cfg.get("ordered_list_style"), index)
    else:
        prefix = _bullet_prefix(cfg.get("list_style"))
    line = f"{prefix} {text}".strip() if prefix else text
    p = doc.add_paragraph(line)
    try:
        indent = float(cfg.get("list_indent_chars") or 2)
        p.paragraph_format.left_indent = Twips(int(indent * 210))
    except (TypeError, ValueError):
        pass


def _set_run_font(run, cell_cfg: dict | None) -> None:
    from docx.shared import Pt  # type: ignore

    if not isinstance(cell_cfg, dict):
        return
    run.font.name = cell_cfg.get("font") or run.font.name
    size = cell_cfg.get("size")
    if size:
        run.font.size = Pt(_pt(str(size), 10.5))
    rgb = _rgb(str(_g(cell_cfg, "text_color", "textColor", default="") or ""))
    if rgb is not None:
        run.font.color.rgb = rgb


def _hex_color(color: str | None, default: str = "000000") -> str:
    """用途：#RGB/#RRGGBB → RRGGBB 大写。"""
    if not color or not isinstance(color, str):
        return default
    s = color.strip().lstrip("#")
    if len(s) == 3:
        s = "".join(c * 2 for c in s)
    if len(s) == 6 and re.fullmatch(r"[0-9a-fA-F]{6}", s):
        return s.upper()
    return default


def _strict_bool(value: Any, default: bool = False) -> bool:
    """用途：把常见布尔值安全归一化，避免字符串 false 被当成真。"""
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"true", "1", "yes", "on"}:
            return True
        if normalized in {"false", "0", "no", "off", ""}:
            return False
    return default


def _heading_border_cfg(template: dict | None) -> dict:
    """用途：读取并清洗标题段落边框配置，兼容 snake_case 与 camelCase。"""
    if not isinstance(template, dict):
        return {
            "enabled": False,
            "min_heading_left_enabled": False,
            "border_color": "000000",
            "level_cell_colors": [],
        }
    raw = _g(template, "heading_border", "headingBorder", default={}) or {}
    if not isinstance(raw, dict):
        raw = {}
    raw_colors = _g(
        raw,
        "level_cell_colors",
        "levelCellColors",
        default=[],
    ) or []
    colors: list[str] = []
    if isinstance(raw_colors, list):
        colors = [
            _hex_color(str(color), "FFFFFF")
            for color in raw_colors[:6]
        ]
    return {
        "enabled": _strict_bool(_g(raw, "enabled", default=False)),
        "min_heading_left_enabled": _strict_bool(
            _g(
                raw,
                "min_heading_left_enabled",
                "minHeadingLeftEnabled",
                default=False,
            )
        ),
        "border_color": _hex_color(
            str(_g(raw, "border_color", "borderColor", default="#000000")),
            "000000",
        ),
        "level_cell_colors": colors,
    }


def _markdown_heading_style_level(hashes: str) -> int:
    """用途：Markdown 井号数映射为 Word Heading 级别（与 write_markdown_body 一致）。"""
    return min(len(hashes) + 1, 4)


def _collect_markdown_heading_levels(text: str) -> list[int]:
    """用途：按文档顺序收集 Markdown 标题的 Word 级别序列。"""
    levels: list[int] = []
    for line in (text or "").replace("\r\n", "\n").split("\n"):
        match = re.match(r"^(#{1,4})\s+(.+)$", line.strip())
        if match:
            levels.append(_markdown_heading_style_level(match.group(1)))
    return levels


def _leaf_flags_from_levels(levels: list[int]) -> list[bool]:
    """
    用途：根据标题级别序列判定每个标题是否为叶子（其后直至同级/更浅标题前无更深标题）。
    对接：大纲 children、Markdown 标题预扫描。
    """
    flags: list[bool] = []
    total = len(levels)
    for index, level in enumerate(levels):
        is_leaf = True
        for later in levels[index + 1 : total]:
            if later <= level:
                break
            if later > level:
                is_leaf = False
                break
        flags.append(is_leaf)
    return flags


def _outline_node_is_leaf(node: dict) -> bool:
    """用途：大纲节点无下级标题子节点时视为叶子。"""
    children = node.get("children")
    if not isinstance(children, list) or not children:
        return True
    return not any(isinstance(child, dict) for child in children)


_PPR_CHILD_ORDER = (
    "pStyle",
    "keepNext",
    "keepLines",
    "pageBreakBefore",
    "framePr",
    "widowControl",
    "numPr",
    "suppressLineNumbers",
    "pBdr",
    "shd",
    "tabs",
    "suppressAutoHyphens",
    "kinsoku",
    "wordWrap",
    "overflowPunct",
    "topLinePunct",
    "autoSpaceDE",
    "autoSpaceDN",
    "bidi",
    "adjustRightInd",
    "snapToGrid",
    "spacing",
    "ind",
    "contextualSpacing",
    "mirrorIndents",
    "suppressOverlap",
    "jc",
    "textDirection",
    "textAlignment",
    "textboxTightWrap",
    "outlineLvl",
    "divId",
    "cnfStyle",
    "rPr",
    "sectPr",
    "pPrChange",
)
_PPR_CHILD_RANK = {name: index for index, name in enumerate(_PPR_CHILD_ORDER)}


def _insert_ppr_child_in_order(paragraph_props, child, local_name: str) -> None:
    """用途：按 OOXML CT_PPr 顺序插入段落属性，兼容严格校验器。"""
    target_rank = _PPR_CHILD_RANK[local_name]
    for index, existing in enumerate(paragraph_props):
        existing_name = existing.tag.rsplit("}", 1)[-1]
        existing_rank = _PPR_CHILD_RANK.get(existing_name, len(_PPR_CHILD_ORDER))
        if existing_rank > target_rank:
            paragraph_props.insert(index, child)
            return
    paragraph_props.append(child)


def _apply_heading_border(
    paragraph,
    cfg: dict | None,
    *,
    level_index: int,
    is_leaf: bool = False,
) -> None:
    """
    用途：为单个标题段落写入四边描边和对应级别底色；
    叶子标题在双开关开启时把左侧边框升为 2.25pt / space=6 的强调线。
    """
    if not isinstance(cfg, dict) or not cfg.get("enabled"):
        return

    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    border_color = _hex_color(str(cfg.get("border_color") or ""), "000000")
    colors = cfg.get("level_cell_colors")
    fill_color = "FFFFFF"
    if isinstance(colors, list) and colors:
        safe_index = max(0, min(int(level_index), len(colors) - 1))
        fill_color = _hex_color(str(colors[safe_index]), "FFFFFF")

    emphasize_left = bool(cfg.get("min_heading_left_enabled")) and bool(is_leaf)

    paragraph_props = paragraph._p.get_or_add_pPr()
    for child in list(paragraph_props):
        if child.tag in {qn("w:pBdr"), qn("w:shd")}:
            paragraph_props.remove(child)

    borders = OxmlElement("w:pBdr")
    for edge_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        if edge_name == "left" and emphasize_left:
            # 2.25 pt = 18 个 1/8 pt；space 为文字与边框间距（磅）
            edge.set(qn("w:sz"), "18")
            edge.set(qn("w:space"), "6")
        else:
            edge.set(qn("w:sz"), "8")
            edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), border_color)
        borders.append(edge)
    _insert_ppr_child_in_order(paragraph_props, borders, "pBdr")

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_color)
    _insert_ppr_child_in_order(paragraph_props, shading, "shd")


def _add_heading(
    doc,
    text: str,
    *,
    level: int,
    heading_border_cfg: dict | None = None,
    is_leaf: bool = False,
):
    """用途：创建 Word 标题，并按实际 Heading 级别与叶子标记应用边框。"""
    paragraph = doc.add_heading(text, level=level)
    if level >= 1:
        _apply_heading_border(
            paragraph,
            heading_border_cfg,
            level_index=min(level - 1, 5),
            is_leaf=is_leaf,
        )
    return paragraph


def _shade_cell(cell, color: str | None) -> None:
    """用途：单元格底色 #RRGGBB。"""
    hex_color = _hex_color(color, "")
    if not hex_color:
        return
    try:
        from docx.oxml.ns import qn  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    except Exception:
        pass


def _set_cell_border(table, border_color: str | None, border_width: float | None) -> None:
    """用途：统一表格边框。"""
    try:
        from docx.oxml.ns import qn, nsdecls  # type: ignore
        from docx.oxml import parse_xml  # type: ignore

        color = _hex_color(border_color, "000000")
        try:
            # border_width 前端多为 pt 量级 0.5~1.5；OOXML sz 单位 1/8 pt
            sz = max(4, int(float(border_width or 0.75) * 8))
        except (TypeError, ValueError):
            sz = 6
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f"<w:tblPr {nsdecls('w')}/>"
        )
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f"</w:tblBorders>"
        )
        # 移除旧 borders
        for child in list(tblPr):
            if child.tag == qn("w:tblBorders"):
                tblPr.remove(child)
        tblPr.append(borders)
    except Exception:
        pass


def add_styled_table(
    doc,
    rows: list[list[str]],
    table_cfg: dict | None = None,
    *,
    has_header: bool = True,
) -> None:
    """
    用途：写入带模板样式的表格。
    rows: 二维字符串；首行可为表头。
    """
    if not rows or not rows[0]:
        return
    cfg = table_cfg if isinstance(table_cfg, dict) else {}
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    header_cfg = _g(cfg, "header_row", "headerRow", default={}) or {}
    first_col_cfg = _g(cfg, "first_column", "firstColumn", default={}) or {}
    body_cfg = _g(cfg, "body_cell", "bodyCell", default={}) or {}
    if not isinstance(header_cfg, dict):
        header_cfg = {}
    if not isinstance(first_col_cfg, dict):
        first_col_cfg = {}
    if not isinstance(body_cfg, dict):
        body_cfg = {}

    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            val = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            is_header = has_header and ri == 0
            is_first_col = ci == 0 and not is_header
            if is_header:
                style = header_cfg
            elif is_first_col and first_col_cfg:
                style = first_col_cfg
            else:
                style = body_cfg
            _set_run_font(run, style if style else None)
            try:
                p.alignment = _align(str(style.get("alignment") or "左对齐"))
            except Exception:
                pass
            bg = str(
                _g(style, "background_color", "backgroundColor", default="") or ""
            )
            if bg:
                _shade_cell(cell, bg)

    _set_cell_border(
        table,
        str(_g(cfg, "border_color", "borderColor", default="#000000") or "#000000"),
        _g(cfg, "border_width", "borderWidth", default=0.75),
    )
    # full_width：尽量拉满
    if _g(cfg, "full_width", "fullWidth", default=True):
        try:
            from docx.shared import Inches, Twips  # type: ignore

            section = doc.sections[0]
            usable = section.page_width - section.left_margin - section.right_margin
            table.autofit = True
            # 设总宽
            tbl = table._tbl
            from docx.oxml.ns import qn  # type: ignore
            from docx.oxml import OxmlElement  # type: ignore

            tblPr = tbl.tblPr
            if tblPr is not None:
                tblW = OxmlElement("w:tblW")
                tblW.set(qn("w:type"), "dxa")
                tblW.set(qn("w:w"), str(int(usable)))
                for child in list(tblPr):
                    if child.tag == qn("w:tblW"):
                        tblPr.remove(child)
                tblPr.append(tblW)
        except Exception:
            pass


_RESPONSE_MATRIX_KIND_LABELS = {
    "requirement": "技术要求",
    "scoring": "评分点",
}
_RESPONSE_MATRIX_STATUS_LABELS = {
    "uncovered": "未覆盖",
    "partial": "部分覆盖",
    "covered": "已覆盖",
    "waived": "不响应",
}


def _index_node_titles(nodes: Any) -> dict[str, str]:
    """用途：递归建立大纲或章节 id 到标题的安全索引。"""
    titles: dict[str, str] = {}
    if not isinstance(nodes, list):
        return titles
    stack = list(nodes)
    while stack:
        node = stack.pop()
        if not isinstance(node, dict):
            continue
        node_id = str(node.get("id") or "").strip()
        title = str(node.get("title") or "").strip()
        if node_id and title:
            titles[node_id] = title
        children = node.get("children")
        if isinstance(children, list):
            stack.extend(children)
    return titles


def _response_matrix_link_labels(
    item: dict[str, Any],
    chapter_titles: dict[str, str],
    outline_titles: dict[str, str],
) -> str:
    """用途：将已收敛的关联 id 显示为可读位置，不泄露内部标识。"""
    labels: list[str] = []
    for chapter_id in item.get("chapterIds") or []:
        title = chapter_titles.get(str(chapter_id))
        if title:
            labels.append(f"正文：{title}")
    for outline_id in item.get("outlineNodeIds") or []:
        title = outline_titles.get(str(outline_id))
        if title:
            labels.append(f"大纲：{title}")
    return "；".join(labels)


def _add_response_matrix_table(
    doc,
    response_matrix: Any,
    outline: Any,
    chapters: Any,
    table_cfg: dict | None = None,
    *,
    heading_border_cfg: dict | None = None,
) -> bool:
    """
    用途：将技术标响应矩阵写入 Word，并在导出时再次过滤失效关联。
    对接：editor_state_service.reconcile_response_matrix；build_docx_bytes 技术标分支。
    """
    from app.services.editor_state_service import reconcile_response_matrix

    rows = reconcile_response_matrix(response_matrix, outline, chapters)
    if not rows:
        return False

    chapter_titles = _index_node_titles(chapters)
    outline_titles = _index_node_titles(outline)
    table_rows: list[list[str]] = [
        ["类型", "要求/评分点", "权重", "响应状态", "关联位置", "备注"]
    ]
    for item in rows:
        kind = str(item.get("kind") or "")
        status = str(item.get("status") or "uncovered")
        table_rows.append(
            [
                _RESPONSE_MATRIX_KIND_LABELS.get(kind, "待确认"),
                str(item.get("sourceText") or ""),
                str(item.get("weight") or ""),
                _RESPONSE_MATRIX_STATUS_LABELS.get(status, "未覆盖"),
                _response_matrix_link_labels(item, chapter_titles, outline_titles),
                str(item.get("notes") or ""),
            ]
        )

    _add_heading(
        doc,
        "六、响应矩阵",
        level=1,
        heading_border_cfg=heading_border_cfg,
    )
    add_styled_table(doc, table_rows, table_cfg, has_header=True)
    return True


def _parse_md_table_lines(lines: list[str]) -> list[list[str]] | None:
    """用途：解析连续 | 表格行；失败返回 None。"""
    if not lines:
        return None
    rows: list[list[str]] = []
    for i, line in enumerate(lines):
        s = line.strip()
        if not s.startswith("|"):
            return None
        # 分隔行 |---|---|
        if re.match(r"^\|[\s:\-|]+\|$", s):
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)
    return rows if rows else None


_PROJECT_IMAGE_LINE_RE = re.compile(
    r'^\s*!\[(?P<alt>[^\]]*)\]\(\s*biaoshu-image://'
    r'(?P<file_id>[^\s)]+)(?:\s+"(?P<caption>[^"]*)")?\s*\)\s*$'
)
_PROJECT_IMAGE_ID_RE = re.compile(r"^file_[0-9a-f]{16}$")


def _add_image_warning(doc, image_warnings: list[str], reason: str) -> None:
    """用途：以可见段落和任务结果同步报告图片降级原因。"""
    message = f"图片引用无效：{reason}"
    doc.add_paragraph(f"【{message}】")
    image_warnings.append(message)


def _add_project_image(doc, image_path: Path, caption: str | None, image_cfg: dict) -> None:
    """用途：按模板版心和图注样式向 Word 插入已受控验证的本地图片。"""
    from docx.shared import Pt  # type: ignore

    section = doc.sections[0]
    available_width = (
        int(section.page_width)
        - int(section.left_margin)
        - int(section.right_margin)
    )
    if available_width <= 0:
        raise ValueError("页面可用宽度无效")
    width = int(available_width * float(image_cfg["max_width_percent"]) / 100)
    paragraph = doc.add_paragraph()
    paragraph.alignment = _align(image_cfg["alignment"])
    paragraph.add_run().add_picture(str(image_path), width=width)

    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = _align(image_cfg["caption_alignment"])
        run = caption_paragraph.add_run(caption)
        run.font.name = image_cfg["caption_font"]
        run.font.size = Pt(_pt(image_cfg["caption_size"], 9))
        run.bold = bool(image_cfg["caption_bold"])
        run.italic = bool(image_cfg["caption_italic"])


def write_markdown_body(
    doc,
    text: str,
    *,
    list_cfg: dict | None = None,
    table_cfg: dict | None = None,
    heading_border_cfg: dict | None = None,
    image_cfg: dict | None = None,
    image_resolver=None,
    image_warnings: list[str] | None = None,
) -> None:
    """
    用途：把章节 Markdown 粗解析为段落/列表/表格写入 docx。
    对接：商务标整包 Markdown、技术标 chapters[].body、heading_border。
    支持：普通段、-/* 无序、1. 有序、| 表格、项目内图片独占行。
    """
    lines = (text or "").replace("\r\n", "\n").split("\n")
    # 写入前预计算叶子标题，禁止流式猜测
    heading_leaf_flags = _leaf_flags_from_levels(
        _collect_markdown_heading_levels(text)
    )
    heading_cursor = 0
    resolved_image_cfg = image_cfg or _image_cfg(None)
    resolved_image_warnings = image_warnings if image_warnings is not None else []
    i = 0
    ol_index = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # 图片只接受项目内受控协议，禁止从 Markdown 请求外链或任意磁盘路径。
        if "biaoshu-image://" in stripped:
            image_match = _PROJECT_IMAGE_LINE_RE.fullmatch(stripped)
            if image_match is None:
                _add_image_warning(doc, resolved_image_warnings, "图片语法不符合约定")
                i += 1
                ol_index = 0
                continue
            file_id = image_match.group("file_id")
            if not _PROJECT_IMAGE_ID_RE.fullmatch(file_id):
                _add_image_warning(doc, resolved_image_warnings, "图片标识不合法")
                i += 1
                ol_index = 0
                continue
            if image_resolver is None:
                _add_image_warning(doc, resolved_image_warnings, "导出服务未配置图片解析器")
                i += 1
                ol_index = 0
                continue
            try:
                image_path = image_resolver(file_id)
                _add_project_image(
                    doc,
                    image_path,
                    image_match.group("caption"),
                    resolved_image_cfg,
                )
            except (KeyError, FileNotFoundError, OSError, ValueError):
                _add_image_warning(doc, resolved_image_warnings, "图片不存在、越权或已损坏")
            i += 1
            ol_index = 0
            continue

        # 表格块
        if stripped.startswith("|"):
            block = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                block.append(lines[j].strip())
                j += 1
            table_rows = _parse_md_table_lines(block)
            if table_rows and len(table_rows) >= 1:
                add_styled_table(doc, table_rows, table_cfg, has_header=True)
                i = j
                ol_index = 0
                continue

        # 无序列表
        m_ul = re.match(r"^[-*+]\s+(.+)$", stripped)
        if m_ul:
            _add_list_item(doc, m_ul.group(1), ordered=False, list_cfg=list_cfg)
            i += 1
            ol_index = 0
            continue

        # 有序列表
        m_ol = re.match(r"^(\d+)[.)、]\s+(.+)$", stripped)
        if m_ol:
            ol_index = int(m_ol.group(1)) if m_ol.group(1) else ol_index + 1
            _add_list_item(
                doc,
                m_ol.group(2),
                ordered=True,
                index=ol_index,
                list_cfg=list_cfg,
            )
            i += 1
            continue

        # 小标题 ##
        m_h = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if m_h:
            level = _markdown_heading_style_level(m_h.group(1))
            is_leaf = True
            if heading_cursor < len(heading_leaf_flags):
                is_leaf = heading_leaf_flags[heading_cursor]
                heading_cursor += 1
            _add_heading(
                doc,
                m_h.group(2).strip(),
                level=level,
                heading_border_cfg=heading_border_cfg,
                is_leaf=is_leaf,
            )
            i += 1
            ol_index = 0
            continue

        doc.add_paragraph(stripped)
        i += 1
        ol_index = 0


def _apply_page_setup(doc, page: dict) -> None:
    """用途：纸张、方向、页边距、页眉页脚、页码。"""
    from docx.shared import Cm, Mm, Pt, Twips  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    section = doc.sections[0]

    paper = str(_g(page, "paper_size", "paperSize", default="a4") or "a4").lower()
    dims = _PAPER_MM.get(paper, _PAPER_MM["a4"])
    orientation = str(_g(page, "orientation", default="portrait") or "portrait").lower()
    w_mm, h_mm = dims
    if orientation == "landscape":
        w_mm, h_mm = h_mm, w_mm
    section.page_width = Mm(w_mm)
    section.page_height = Mm(h_mm)

    def margin(key_cm: str, key_alt: str, default_cm: float):
        v = _g(page, key_cm, key_alt)
        try:
            return Cm(float(v))
        except (TypeError, ValueError):
            return Cm(default_cm)

    section.top_margin = margin("margin_top_cm", "top", 2.54)
    section.bottom_margin = margin("margin_bottom_cm", "bottom", 2.54)
    section.left_margin = margin("margin_left_cm", "left", 3.17)
    section.right_margin = margin("margin_right_cm", "right", 3.17)

    first_diff = bool(
        _g(page, "first_page_different", "firstPageDifferent", default=False)
    )
    section.different_first_page_header_footer = first_diff

    # 页眉
    if _g(page, "header_enabled", "headerEnabled", default=False):
        text = str(_g(page, "header_text", "headerText", default="") or "")
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.name = str(_g(page, "header_font", "headerFont", default="宋体") or "宋体")
        run.font.size = Pt(_pt(str(_g(page, "header_size", "headerSize", default="小五")), 9))
        rgb = _rgb(str(_g(page, "header_color", "headerColor", default="") or ""))
        if rgb is not None:
            run.font.color.rgb = rgb
        p.alignment = _align(str(_g(page, "header_alignment", "headerAlignment", default="居中") or ""))

    # 页脚 + 页码
    footer_on = bool(_g(page, "footer_enabled", "footerEnabled", default=False))
    page_num_on = bool(_g(page, "page_number_enabled", "pageNumberEnabled", default=False))
    if footer_on or page_num_on:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = _align(
            str(_g(page, "footer_alignment", "footerAlignment", default="居中") or "居中")
        )
        font_name = str(_g(page, "footer_font", "footerFont", default="宋体") or "宋体")
        font_size = Pt(_pt(str(_g(page, "footer_size", "footerSize", default="小五")), 9))
        rgb = _rgb(str(_g(page, "footer_color", "footerColor", default="") or ""))

        if footer_on:
            ftext = str(_g(page, "footer_text", "footerText", default="") or "")
            if ftext:
                r = p.add_run(ftext + ("　" if page_num_on else ""))
                r.font.name = font_name
                r.font.size = font_size
                if rgb is not None:
                    r.font.color.rgb = rgb

        if page_num_on:
            fmt = str(
                _g(page, "page_number_format", "pageNumberFormat", default="第{page}页")
                or "第{page}页"
            )
            # 拆成「前缀 + PAGE 域 + 后缀」
            if "{page}" in fmt:
                before, after = fmt.split("{page}", 1)
            else:
                before, after = fmt, ""

            def add_run_txt(s: str) -> None:
                if not s:
                    return
                r = p.add_run(s)
                r.font.name = font_name
                r.font.size = font_size
                if rgb is not None:
                    r.font.color.rgb = rgb

            add_run_txt(before)
            # PAGE 域
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run_el = p.add_run()._r
            run_el.append(fld_begin)
            run_el.append(instr)
            run_el.append(fld_end)
            add_run_txt(after)

            start = _g(page, "page_number_start", "pageNumberStart", default=1)
            try:
                start_i = int(start)
            except (TypeError, ValueError):
                start_i = 1
            if start_i != 1:
                try:
                    section.start_at = start_i  # type: ignore[attr-defined]
                except Exception:
                    pass

        dist = _g(page, "footer_distance_cm", "footerDistanceCm")
        try:
            if dist is not None:
                section.footer_distance = Cm(float(dist))
        except (TypeError, ValueError):
            pass


def _apply_template_styles(doc, cfg: dict) -> None:
    """用途：把 ExportFormatConfig 字段应用到 docx 样式与节属性。"""
    from docx.shared import Pt, Twips  # type: ignore

    page = _g(cfg, "page", "page_setup", default={}) or {}
    if isinstance(page, dict):
        try:
            _apply_page_setup(doc, page)
        except Exception:
            # 页设置失败不阻断导出
            pass

    body = _g(cfg, "body_text", "bodyText", default={}) or {}
    if isinstance(body, dict):
        style = doc.styles["Normal"]
        font = style.font
        font.name = body.get("font") or "宋体"
        font.size = Pt(_pt(body.get("size"), 12))
        pf = style.paragraph_format
        try:
            pf.line_spacing = float(
                _g(body, "line_spacing_multiple", "lineSpacingMultiple", default=1.5)
                or 1.5
            )
        except (TypeError, ValueError):
            pf.line_spacing = 1.5
        try:
            sb = float(_g(body, "spacing_before_pt", "spacingBeforePt", default=0) or 0)
            sa = float(_g(body, "spacing_after_pt", "spacingAfterPt", default=0) or 0)
            pf.space_before = Pt(sb)
            pf.space_after = Pt(sa)
        except (TypeError, ValueError):
            pass
        try:
            pf.alignment = _align(str(body.get("alignment") or ""))
        except Exception:
            pass
        indent = _g(body, "first_line_indent_chars", "firstLineIndentChars", default=0) or 0
        try:
            # 约 1 字符 ≈ 210 twips（小四）
            pf.first_line_indent = Twips(int(float(indent) * 210))
        except (TypeError, ValueError):
            pass

    headings = _g(cfg, "headings", default=[]) or []
    if isinstance(headings, list):
        for i, h in enumerate(headings[:6]):
            if not isinstance(h, dict):
                continue
            style_name = f"Heading {i + 1}"
            try:
                st = doc.styles[style_name]
            except KeyError:
                continue
            st.font.name = h.get("font") or "黑体"
            st.font.size = Pt(_pt(h.get("size"), max(10, 16 - i * 2)))
            st.font.bold = bool(h.get("bold", True))
            rgb = _rgb(str(_g(h, "text_color", "textColor", default="") or ""))
            if rgb is not None:
                st.font.color.rgb = rgb
            try:
                st.paragraph_format.alignment = _align(h.get("alignment"))
            except Exception:
                pass
            try:
                sb = float(
                    _g(h, "spacing_before_pt", "spacingBeforePt", default=12) or 12
                )
                sa = float(
                    _g(h, "spacing_after_pt", "spacingAfterPt", default=6) or 6
                )
                st.paragraph_format.space_before = Pt(sb)
                st.paragraph_format.space_after = Pt(sa)
            except (TypeError, ValueError):
                pass
            try:
                ls = float(_g(h, "line_spacing", "lineSpacing", default=1.0) or 1.0)
                st.paragraph_format.line_spacing = ls
            except (TypeError, ValueError):
                pass


# Windows 保留设备名（整名大小写不敏感）；人读文件名若命中则尾加 _
_WINDOWS_RESERVED_BASENAMES = frozenset(
    {
        "CON",
        "PRN",
        "AUX",
        "NUL",
        *{f"COM{i}" for i in range(1, 10)},
        *{f"LPT{i}" for i in range(1, 10)},
    }
)
# Windows 非法路径字符
_WINDOWS_ILLEGAL_FILENAME_CHARS = frozenset('<>:"/\\|?*')


def build_safe_docx_filename(project_name: str | None) -> str:
    """
    用途：唯一安全人读 DOCX 文件名收敛；任务 result.filename 与下载 Content-Disposition 共用。
    规则：去控制字符与 <>:"/\\|?* → 去首尾空白/尾点空格 → 循环去重复 .docx →
          基础名最多 100 码点 → 保留名尾加 _ → 空回退「标书」→ 单次 .docx。
    对接：build_docx_bytes 返回值；export 下载路由 FileResponse.filename。
    二次开发：不得用手拼未转义头；磁盘/URL 仍用 export_*.docx 随机 storedName。
    """
    raw = project_name if isinstance(project_name, str) else ""
    cleaned_chars: list[str] = []
    for ch in raw:
        code = ord(ch)
        # 移除 C0 控制字符、DEL 与 C1（U+007F–U+009F，含 TAB/NEL）；不扩大到 Cf 等
        if code < 32 or 127 <= code <= 159:
            continue
        if ch in _WINDOWS_ILLEGAL_FILENAME_CHARS:
            continue
        cleaned_chars.append(ch)
    base = "".join(cleaned_chars).strip().rstrip(" .")
    # 循环剥除重复 .docx（大小写不敏感），并收敛尾点/空格
    while base.lower().endswith(".docx"):
        base = base[: -len(".docx")].rstrip(" .")
    if len(base) > 100:
        base = base[:100].rstrip(" .")
    if not base:
        base = "标书"
    if base.upper() in _WINDOWS_RESERVED_BASENAMES:
        base = f"{base}_"
    return f"{base}.docx"


def build_docx_bytes(
    db: Session,
    workspace_id: str,
    project_id: str,
    *,
    mode: str | None = None,
    image_warnings: list[str] | None = None,
) -> tuple[bytes, str]:
    """
    用途：生成 Word 文档（封面 + 正文及技术标响应矩阵，应用默认导出模板）。
    对接：export 同步/异步任务；settings.export_format_json。
    mode：technical（默认）| business（商务标册，读 business_json）。
    """
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，无法导出 Word") from exc

    project = get_project(db, workspace_id, project_id)
    state = db.get(ProjectEditorStateRow, project_id)
    template = settings_service.get_export_format(db, workspace_id)

    export_mode = (mode or "").strip().lower()
    if not export_mode and getattr(project, "kind", None) == "business":
        export_mode = "business"
    if export_mode not in ("business", "technical"):
        export_mode = "technical"

    doc = Document()
    h1_page_break = False
    headings_cfg: list = []
    list_cfg = _body_list_cfg(template)
    table_cfg = _table_cfg(template)
    heading_border_cfg = _heading_border_cfg(template)
    image_cfg = _image_cfg(template)
    resolved_image_warnings = image_warnings if image_warnings is not None else []
    settings = get_settings()

    def resolve_project_image(file_id: str) -> Path:
        _, image_path = file_service.resolve_project_image(
            db,
            workspace_id,
            project_id,
            settings,
            file_id,
        )
        return image_path

    if template:
        try:
            _apply_template_styles(doc, template)
            h1_page_break = bool(
                _g(
                    template,
                    "heading_level1_page_break_before",
                    "headingLevel1PageBreakBefore",
                    default=False,
                )
            )
            raw_h = _g(template, "headings", default=[]) or []
            if isinstance(raw_h, list):
                headings_cfg = raw_h
        except Exception:
            pass

    is_business = export_mode == "business"
    title = project.name or ("商务标" if is_business else "技术标")
    cover_label = "商务标书" if is_business else "技术标书"

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(cover_label)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tpl_name = ""
    if template:
        tpl_name = str(
            _g(template, "template_name", "name", "templateName", default="") or ""
        )
    meta_text = (
        f"行业：{project.industry or '通用'}　｜　"
        f"导出日期：{datetime.now().strftime('%Y-%m-%d')}　｜　"
        f"状态：{project.status}"
        + (f"　｜　模板：{tpl_name}" if tpl_name else "")
    )
    mr = meta.add_run(meta_text)
    mr.font.size = Pt(10.5)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_page_break()

    doc.add_heading(title, level=0)

    # —— 商务标：整包 Markdown 导出 ——
    if is_business:
        from app.services import business_task_service, editor_state_service

        ed = editor_state_service.get_editor_state(db, workspace_id, project_id)
        md = business_task_service.build_business_markdown(ed, title)
        write_markdown_body(
            doc,
            md,
            list_cfg=list_cfg,
            table_cfg=table_cfg,
            heading_border_cfg=heading_border_cfg,
            image_cfg=image_cfg,
            image_resolver=resolve_project_image,
            image_warnings=resolved_image_warnings,
        )
        buf = io.BytesIO()
        doc.save(buf)
        # 人读名仅由权威 project.name 收敛；封面/正文 title 保持原样
        filename = build_safe_docx_filename(project.name)
        return buf.getvalue(), filename

    overview = (state.analysis_overview if state else None) or ""
    analysis = _loads(state.analysis_json) if state else None
    if isinstance(analysis, dict) and analysis.get("overview"):
        overview = analysis["overview"]

    if overview.strip():
        _add_heading(
            doc,
            "一、项目概述 / 招标分析",
            level=1,
            heading_border_cfg=heading_border_cfg,
        )
        for para in overview.strip().split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    if isinstance(analysis, dict):
        tr = analysis.get("techRequirements") or []
        if tr:
            _add_heading(
                doc,
                "技术要求",
                level=2,
                heading_border_cfg=heading_border_cfg,
            )
            for item in tr:
                _add_list_item(doc, str(item), ordered=False, list_cfg=list_cfg)
        sp = analysis.get("scoringPoints") or []
        if sp:
            _add_heading(
                doc,
                "评分点",
                level=2,
                heading_border_cfg=heading_border_cfg,
            )
            # 有权重时用表格；否则有序列表
            table_rows: list[list[str]] = [["评分项", "权重"]]
            use_table = False
            for p in sp:
                if isinstance(p, dict):
                    name = str(p.get("name") or "")
                    weight = str(p.get("weight") or "")
                    if weight:
                        use_table = True
                    table_rows.append([name, weight])
                else:
                    table_rows.append([str(p), ""])
            if use_table and len(table_rows) > 1:
                add_styled_table(doc, table_rows, table_cfg, has_header=True)
            else:
                for idx, p in enumerate(sp, 1):
                    if isinstance(p, dict):
                        label = f"{p.get('name', '')}　{p.get('weight', '')}".strip()
                    else:
                        label = str(p)
                    _add_list_item(
                        doc, label, ordered=True, index=idx, list_cfg=list_cfg
                    )
        rr = analysis.get("rejectionRisks") or []
        if rr:
            _add_heading(
                doc,
                "废标风险",
                level=2,
                heading_border_cfg=heading_border_cfg,
            )
            for item in rr:
                _add_list_item(doc, str(item), ordered=False, list_cfg=list_cfg)

    parsed = (state.parsed_markdown if state else None) or ""
    if parsed.strip():
        _add_heading(
            doc,
            "二、招标文件解析摘录",
            level=1,
            heading_border_cfg=heading_border_cfg,
        )
        clip = parsed.strip()
        if len(clip) > 20000:
            clip = clip[:20000] + "\n\n…（导出已截断）"
        for para in clip.split("\n"):
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(10.5)

    outline = _loads(state.outline_json) if state else None
    if isinstance(outline, list) and outline:
        _add_heading(
            doc,
            "三、目录大纲",
            level=1,
            heading_border_cfg=heading_border_cfg,
            is_leaf=False,
        )
        outline_num = HeadingNumberer(headings_cfg)

        def walk(nodes: list, depth: int = 1) -> None:
            # depth: 1=一级节点 → numbering level 0
            for n in nodes:
                if not isinstance(n, dict):
                    continue
                raw_title = str(n.get("title") or "未命名")
                num_level = max(0, min(depth - 1, 5))
                prefix = outline_num.next_prefix(num_level)
                t = compose_heading_text(raw_title, prefix)
                # Word 样式：大纲挂在「三、」下，用 Heading 2+ 更合适
                style_level = min(depth + 1, 4)
                _add_heading(
                    doc,
                    t,
                    level=style_level,
                    heading_border_cfg=heading_border_cfg,
                    is_leaf=_outline_node_is_leaf(n),
                )
                desc = n.get("description")
                if desc:
                    doc.add_paragraph(str(desc))
                children = n.get("children")
                if isinstance(children, list) and children:
                    walk(children, depth + 1)

        walk(outline)

    chapters = _loads(state.chapters_json) if state else None
    if isinstance(chapters, list) and chapters:
        _add_heading(
            doc,
            "四、正文",
            level=1,
            heading_border_cfg=heading_border_cfg,
            is_leaf=False,
        )
        chapter_num = HeadingNumberer(headings_cfg)
        first_ch = True
        for ch in chapters:
            if not isinstance(ch, dict):
                continue
            if h1_page_break and not first_ch:
                doc.add_page_break()
            first_ch = False
            raw_title = str(ch.get("title") or "章节")
            # 正文一级章 → headings[0]
            prefix = chapter_num.next_prefix(0)
            ch_title = compose_heading_text(raw_title, prefix)
            _add_heading(
                doc,
                ch_title,
                level=2,
                heading_border_cfg=heading_border_cfg,
            )
            body = str(ch.get("body") or "").strip()
            if not body:
                doc.add_paragraph("（本章暂无正文）")
                continue
            write_markdown_body(
                doc,
                body,
                list_cfg=list_cfg,
                table_cfg=table_cfg,
                heading_border_cfg=heading_border_cfg,
                image_cfg=image_cfg,
                image_resolver=resolve_project_image,
                image_warnings=resolved_image_warnings,
            )

    facts = _loads(state.facts_json) if state else None
    if isinstance(facts, list) and facts:
        _add_heading(
            doc,
            "五、全局事实",
            level=1,
            heading_border_cfg=heading_border_cfg,
        )
        for f in facts:
            if not isinstance(f, dict):
                continue
            cat = f.get("category") or ""
            content = f.get("content") or ""
            _add_list_item(
                doc, f"[{cat}] {content}", ordered=False, list_cfg=list_cfg
            )

    _add_response_matrix_table(
        doc,
        _loads(getattr(state, "response_matrix_json", None)) if state else None,
        outline,
        chapters,
        table_cfg,
        heading_border_cfg=heading_border_cfg,
    )

    buf = io.BytesIO()
    doc.save(buf)
    # 人读名仅由权威 project.name 收敛；封面/正文 title 保持原样
    filename = build_safe_docx_filename(project.name)
    return buf.getvalue(), filename
